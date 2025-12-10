#!/usr/bin/env python
"""
OSINT Investigation Framework (OIF)
====================================

A comprehensive Open Source Intelligence gathering and analysis platform
for conducting investigations across multiple data sources.

Features:
- Multi-source data ingestion (logs, weather, network, files)
- Pattern recognition and correlation
- Timeline reconstruction
- Entity extraction and relationship mapping
- Automated reporting with evidence chain
- Configurable investigation profiles

Author: OSINT Framework Team
Version: 1.0.0
"""

from __future__ import annotations

import argparse
import bz2
import contextlib
import csv
import datetime
import hashlib
import json
import logging
import os
import re
import sqlite3
import subprocess
import sys
import tempfile
import threading
import uuid
from abc import ABC, abstractmethod
from collections import defaultdict, Counter
from collections.abc import Iterable, Iterator, Callable
from dataclasses import dataclass, field, asdict
from enum import Enum, auto
from functools import wraps
from pathlib import Path
from typing import (
    Any, Dict, List, Optional, Set, Tuple, Union, 
    NamedTuple, TypeVar, Generic, TextIO, Protocol, Generator
)
import urllib.parse
import urllib.request

# External dependencies (install via: pip install requests)
try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

# PDF parsing support
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Excel/Spreadsheet support
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import xlrd
    XLRD_AVAILABLE = True
except ImportError:
    XLRD_AVAILABLE = False

# Email parsing support
try:
    import email
    from email import policy
    from email.parser import BytesParser
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False

# Image processing and metadata extraction
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# OCR support for images
try:
    import pytesseract
    # Configure Tesseract OCR path for Windows
    TESSERACT_PATH = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if Path(TESSERACT_PATH).exists():
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

# Document processing (Word docs)
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# YAML support
try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

# XML parsing (stdlib but wrapped for consistency)
try:
    import xml.etree.ElementTree as ET
    XML_AVAILABLE = True
except ImportError:
    XML_AVAILABLE = False

# Watchdog for file system monitoring
try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler, FileCreatedEvent, FileModifiedEvent
    WATCHDOG_AVAILABLE = True
except ImportError:
    WATCHDOG_AVAILABLE = False

# LangChain integration for enhanced LLM processing
LANGCHAIN_AVAILABLE = False
try:
    from langchain_community.llms import Ollama
    from langchain.prompts import PromptTemplate
    from langchain.chains import LLMChain
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain.callbacks.manager import CallbackManager
    from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
    LANGCHAIN_AVAILABLE = True
except ImportError:
    pass

# =============================================================================
# SECURITY: LOCAL-ONLY MODE
# =============================================================================
# This framework operates entirely locally. No external network calls are made.
# All LLM processing is done via local Ollama instance (localhost:11434).
# No data is sent to external servers, APIs, or cloud services.
LOCAL_ONLY_MODE = True
ALLOWED_HOSTS = ["localhost", "127.0.0.1"]

# =============================================================================
# BASE DIRECTORIES
# =============================================================================

# Base directory for all investigations - all output will be saved here
BASE_INVESTIGATIONS_DIR = Path(__file__).parent / "INVESTIGATIONS"
BASE_INVESTIGATIONS_DIR.mkdir(parents=True, exist_ok=True)

# =============================================================================
# TIMEZONE CONFIGURATION
# =============================================================================
# Modify these settings to change the default timezone for reports
# CST (Central Standard Time) is UTC-6, CDT (Central Daylight Time) is UTC-5
# Set TIMEZONE_OFFSET to the hours offset from UTC (negative for west of UTC)
# Set TIMEZONE_NAME to the display name for the timezone

TIMEZONE_OFFSET = -6  # CST is UTC-6 (change to -5 for CDT)
TIMEZONE_NAME = "CST"  # Display name for timezone
USE_12_HOUR_FORMAT = True  # Set to False for 24-hour time format

def get_current_time() -> datetime.datetime:
    """Get current time adjusted to configured timezone."""
    utc_now = datetime.datetime.utcnow()
    return utc_now + datetime.timedelta(hours=TIMEZONE_OFFSET)

def format_timestamp(dt: Optional[datetime.datetime] = None, include_date: bool = True) -> str:
    """Format a datetime object using configured timezone and format settings.
    
    Args:
        dt: Datetime to format. If None, uses current time.
        include_date: Whether to include the date portion.
    
    Returns:
        Formatted timestamp string.
    """
    if dt is None:
        dt = get_current_time()
    
    if USE_12_HOUR_FORMAT:
        time_format = "%I:%M:%S %p"  # 12-hour with AM/PM
    else:
        time_format = "%H:%M:%S"  # 24-hour
    
    if include_date:
        return f"{dt.strftime('%Y-%m-%d')} {dt.strftime(time_format)} {TIMEZONE_NAME}"
    else:
        return f"{dt.strftime(time_format)} {TIMEZONE_NAME}"

def format_timestamp_compact(dt: Optional[datetime.datetime] = None) -> str:
    """Format a datetime object in compact format for filenames.
    
    Args:
        dt: Datetime to format. If None, uses current time.
    
    Returns:
        Compact timestamp string (YYYYMMDD_HHMMSS).
    """
    if dt is None:
        dt = get_current_time()
    return dt.strftime('%Y%m%d_%H%M%S')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s in %(module)s: %(message)s'
)
logger = logging.getLogger(__name__)

# =============================================================================
# CORE DATA MODELS
# =============================================================================

class InvestigationType(Enum):
    """Types of OSINT investigations supported."""
    PERSON = auto()
    ORGANIZATION = auto()
    DOMAIN = auto()
    IP_ADDRESS = auto()
    EMAIL = auto()
    PHONE = auto()
    SOCIAL_MEDIA = auto()
    CRYPTOCURRENCY = auto()
    VEHICLE = auto()
    LOCATION = auto()
    INCIDENT = auto()
    NETWORK = auto()
    MALWARE = auto()


class DataSourceType(Enum):
    """Types of data sources."""
    LOG_FILE = auto()
    CSV_FILE = auto()
    JSON_FILE = auto()
    API_RESPONSE = auto()
    DATABASE = auto()
    NETWORK_CAPTURE = auto()
    SOCIAL_MEDIA = auto()
    PUBLIC_RECORDS = auto()
    WEATHER_DATA = auto()
    GEOLOCATION = auto()


class SeverityLevel(Enum):
    """Severity levels for findings."""
    CRITICAL = 5
    HIGH = 4
    MEDIUM = 3
    LOW = 2
    INFO = 1


class EntityType(Enum):
    """Types of entities that can be extracted."""
    PERSON = auto()
    EMAIL = auto()
    PHONE = auto()
    IP_ADDRESS = auto()
    URL = auto()
    DOMAIN = auto()
    HASH = auto()
    CRYPTOCURRENCY_WALLET = auto()
    SOCIAL_HANDLE = auto()
    GEOLOCATION = auto()
    TIMESTAMP = auto()
    FILE_PATH = auto()
    MAC_ADDRESS = auto()


# =============================================================================
# NAMED TUPLES FOR STRUCTURED DATA
# =============================================================================

class RawLog(NamedTuple):
    """Raw log entry structure."""
    date: str
    level: str
    module: str
    message: str


class DatedLog(NamedTuple):
    """Log entry with parsed datetime."""
    date: datetime.datetime
    level: str
    module: str
    message: str


class NetworkConnection(NamedTuple):
    """Network connection record."""
    timestamp: datetime.datetime
    source_ip: str
    source_port: int
    dest_ip: str
    dest_port: int
    protocol: str
    bytes_sent: int
    bytes_received: int


class GeoLocation(NamedTuple):
    """Geographic location data."""
    latitude: float
    longitude: float
    accuracy: float
    timestamp: datetime.datetime
    source: str


class WeatherData(NamedTuple):
    """Weather data structure."""
    id: str
    location: str
    max_temp: int
    min_temp: int
    precipitation: float
    timestamp: datetime.datetime


class FileMetadata(NamedTuple):
    """File metadata for evidence tracking."""
    path: Path
    hash_md5: str
    hash_sha256: str
    size: int
    created: datetime.datetime
    modified: datetime.datetime
    accessed: datetime.datetime


class Entity(NamedTuple):
    """Extracted entity from data."""
    type: EntityType
    value: str
    confidence: float
    source: str
    context: str
    timestamp: Optional[datetime.datetime]


class Relationship(NamedTuple):
    """Relationship between two entities."""
    source_entity: Entity
    target_entity: Entity
    relationship_type: str
    confidence: float
    evidence: List[str]


class Finding(NamedTuple):
    """Investigation finding."""
    id: str
    severity: SeverityLevel
    title: str
    description: str
    entities: List[Entity]
    evidence: List[str]
    timestamp: datetime.datetime
    recommendations: List[str]


class TimelineEvent(NamedTuple):
    """Event for timeline reconstruction."""
    timestamp: datetime.datetime
    event_type: str
    description: str
    entities: List[Entity]
    source: str
    confidence: float


# =============================================================================
# CONFIGURATION MANAGEMENT
# =============================================================================

@dataclass
class InvestigationConfig:
    """Configuration for an investigation."""
    name: str
    investigation_type: InvestigationType
    targets: List[str]
    data_sources: List[Path]
    output_dir: Path
    api_keys: Dict[str, str] = field(default_factory=dict)
    custom_patterns: Dict[str, str] = field(default_factory=dict)
    max_depth: int = 3
    timeout: int = 30
    parallel_workers: int = 4
    enable_caching: bool = True
    cache_ttl: int = 3600
    report_format: str = "markdown"
    
    @classmethod
    def from_file(cls, config_path: Path) -> 'InvestigationConfig':
        """Load configuration from file."""
        if config_path.suffix == '.json':
            with config_path.open() as f:
                data = json.load(f)
        elif config_path.suffix == '.py':
            data = cls._load_python_config(config_path)
        else:
            raise ValueError(f"Unsupported config format: {config_path.suffix}")
        
        # Determine output directory - use INVESTIGATIONS base if not specified
        output_dir_str = data.get('output_dir', '')
        if output_dir_str:
            output_dir = Path(output_dir_str)
        else:
            inv_name = data.get('name', 'Unnamed_Investigation').replace(' ', '_').lower()
            output_dir = BASE_INVESTIGATIONS_DIR / inv_name
        
        return cls(
            name=data.get('name', 'Unnamed Investigation'),
            investigation_type=InvestigationType[data.get('type', 'INCIDENT').upper()],
            targets=data.get('targets', []),
            data_sources=[Path(p) for p in data.get('data_sources', [])],
            output_dir=output_dir,
            api_keys=data.get('api_keys', {}),
            custom_patterns=data.get('custom_patterns', {}),
            max_depth=data.get('max_depth', 3),
            timeout=data.get('timeout', 30),
            parallel_workers=data.get('parallel_workers', 4),
            enable_caching=data.get('enable_caching', True),
            cache_ttl=data.get('cache_ttl', 3600),
            report_format=data.get('report_format', 'markdown')
        )
    
    @staticmethod
    def _load_python_config(config_path: Path) -> Dict[str, Any]:
        """Load configuration from Python file."""
        code = compile(
            config_path.read_text(),
            config_path.name,
            "exec"
        )
        locals_dict: Dict[str, Any] = {}
        exec(code, {"__builtins__": __builtins__}, locals_dict)
        return locals_dict.get('config', {})
    
    def to_file(self, config_path: Path) -> None:
        """Save configuration to file."""
        data = {
            'name': self.name,
            'type': self.investigation_type.name,
            'targets': self.targets,
            'data_sources': [str(p) for p in self.data_sources],
            'output_dir': str(self.output_dir),
            'api_keys': self.api_keys,
            'custom_patterns': self.custom_patterns,
            'max_depth': self.max_depth,
            'timeout': self.timeout,
            'parallel_workers': self.parallel_workers,
            'enable_caching': self.enable_caching,
            'cache_ttl': self.cache_ttl,
            'report_format': self.report_format
        }
        with config_path.open('w') as f:
            json.dump(data, f, indent=2)


# =============================================================================
# CUSTOM JSON ENCODER
# =============================================================================

class OSINTJSONEncoder(json.JSONEncoder):
    """Custom JSON encoder that handles Path objects, datetime, and other types."""
    
    def default(self, o):
        if isinstance(o, Path):
            return str(o)
        if isinstance(o, datetime.datetime):
            return o.isoformat()
        if isinstance(o, datetime.date):
            return o.isoformat()
        if isinstance(o, bytes):
            return o.decode('utf-8', errors='ignore')
        if isinstance(o, Enum):
            return o.name
        if hasattr(o, '_asdict'):  # NamedTuple
            return o._asdict()
        if hasattr(o, '__dict__'):
            return o.__dict__
        return super().default(o)


def safe_json_dumps(obj: Any, **kwargs) -> str:
    """Safely dump object to JSON string, handling non-serializable types."""
    return json.dumps(obj, cls=OSINTJSONEncoder, **kwargs)


# =============================================================================
# PATTERN MATCHING ENGINE
# =============================================================================

class PatternLibrary:
    """Library of regex patterns for entity extraction."""
    
    # Core patterns - improved to reduce false positives
    PATTERNS = {
        EntityType.EMAIL: re.compile(
            r'(?P<email>[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            re.IGNORECASE
        ),
        # IP addresses - exclude version numbers by not matching after "version" or software names
        # Match: private ranges, localhost, and public IPs with at least one octet >= 10
        EntityType.IP_ADDRESS: re.compile(
            r'(?<!version\s)(?<!Version\s)(?<![vV]\s)(?<![vV])(?<![0-9a-fA-F.:])(?<!\.\d\.)(?P<ip>'
            r'(?:'
            r'(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2})\.(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.){2}(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9]))|'  # First octet >= 100
            r'(?:(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.){2}(?:25[0-5]|2[0-4][0-9]|1[0-9]{2})\.(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9]))|'  # Third octet >= 100
            r'(?:(?:[1-9][0-9])\.(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.){2}(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9][0-9]?))|'  # First octet 10-99, last octet >= 10
            r'(?:10\.(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.){2}(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9]))|'  # 10.x.x.x private
            r'(?:172\.(?:1[6-9]|2[0-9]|3[01])\.(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.)?(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9]))|'  # 172.16-31.x.x private
            r'(?:192\.168\.(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.)?(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9]))|'  # 192.168.x.x private
            r'(?:127\.0\.0\.[0-9]{1,3})|'  # Localhost
            r'(?:[8-9]\.(?:(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9])\.){2}(?:25[0-5]|2[0-4][0-9]|1[0-9]{2}|[1-9]?[0-9]))'  # 8.x.x.x or 9.x.x.x public
            r')'
            r')(?![0-9.])'
        ),
        EntityType.URL: re.compile(
            r'(?P<url>https?://[^\s<>"{}|\\^`\[\]]+)',
            re.IGNORECASE
        ),
        # Domain - exclude common false positives and require valid TLD
        EntityType.DOMAIN: re.compile(
            r'(?<![a-zA-Z0-9@./_\\-])(?P<domain>'
            r'(?:[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?\.)+(?:com|org|net|edu|gov|mil|io|co|info|biz|us|uk|ca|de|fr|au|ru|cn|in|br|jp|it|nl|se|no|es|pl|ch|at|be|dk|fi|ie|nz|za|mx|ar|cl|kr|tw|hk|sg|my|th|ph|vn|id|pk|bd|ng|eg|ke|ae|sa|il|tr|gr|cz|hu|ro|ua|by|kz|uz|pt|sk|si|hr|rs|bg|ee|lt|lv))'
            r'(?![a-zA-Z0-9])'
        ),
        # Phone - require proper formatting, exclude pure digit sequences that could be hashes
        # Must have separators OR start with + OR have parentheses for area code
        EntityType.PHONE: re.compile(
            r'(?<![0-9a-fA-F])(?P<phone>'
            r'(?:\+1[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4})|'  # +1 format
            r'(?:\(?[0-9]{3}\)[-.\s][0-9]{3}[-.\s]?[0-9]{4})|'  # (xxx) xxx-xxxx with required separator
            r'(?:[0-9]{3}[-.\s][0-9]{3}[-.\s][0-9]{4})'  # xxx-xxx-xxxx with separators
            r')(?![0-9a-fA-F])'
        ),
        # Hash - must be standalone, not part of larger hex string
        EntityType.HASH: re.compile(
            r'(?<![a-fA-F0-9])(?P<hash>'
            r'(?:[a-fA-F0-9]{32}|[a-fA-F0-9]{40}|[a-fA-F0-9]{64})'
            r')(?![a-fA-F0-9])'
        ),
        # Crypto wallet - Bitcoin must start with 1, 3, or bc1; Ethereum with 0x
        # More strict length requirements to avoid matching hash fragments
        EntityType.CRYPTOCURRENCY_WALLET: re.compile(
            r'(?<![a-zA-Z0-9])(?P<wallet>'
            r'(?:bc1[a-zA-HJ-NP-Z0-9]{39,59})|'  # Bech32 Bitcoin (42-62 chars total)
            r'(?:[13][a-km-zA-HJ-NP-Z1-9]{25,34})|'  # Legacy Bitcoin (26-35 chars, Base58)
            r'(?:0x[a-fA-F0-9]{40})'  # Ethereum (42 chars exactly)
            r')(?![a-zA-Z0-9])'
        ),
        # Social handle - must be preceded by whitespace or start of string
        EntityType.SOCIAL_HANDLE: re.compile(
            r'(?:^|(?<=\s))(?P<handle>@[a-zA-Z][a-zA-Z0-9_]{0,14})(?=\s|$|[.,;:!?])',
            re.MULTILINE
        ),
        EntityType.MAC_ADDRESS: re.compile(
            r'(?P<mac>(?:[0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2})'
        ),
        # File path - more strict to avoid matching random text
        EntityType.FILE_PATH: re.compile(
            r'(?P<path>'
            r'(?:[A-Z]:\\(?:[a-zA-Z0-9._\s-]+\\)*[a-zA-Z0-9._\s-]+)|'  # Windows path
            r'(?:/(?:home|usr|var|etc|opt|tmp|root|mnt|media|srv)/[a-zA-Z0-9._/-]+)'  # Unix path with known root
            r')'
        ),
        EntityType.TIMESTAMP: re.compile(
            r'(?P<timestamp>\d{4}-\d{2}-\d{2}[T\s]\d{2}:\d{2}:\d{2}(?:\.\d+)?(?:Z|[+-]\d{2}:?\d{2})?)'
        ),
    }
    
    # Log parsing pattern
    LOG_PATTERN = re.compile(
        r"\[(?P<date>.*?)\]\s+"
        r"(?P<level>\w+)\s+"
        r"in\s+(?P<module>.+?)"
        r":\s+(?P<message>.+)"
    )
    
    # Weather data pattern
    WEATHER_PATTERN = re.compile(
        r'(?P<units>\w+)'  # units, alphanumeric characters
    )
    
    @classmethod
    def extract_entities(
        cls,
        text: str,
        source: str = "unknown",
        custom_patterns: Optional[Dict[str, str]] = None
    ) -> List[Entity]:
        """Extract all entities from text."""
        entities = []
        
        # Apply standard patterns
        for entity_type, pattern in cls.PATTERNS.items():
            for match in pattern.finditer(text):
                entity = Entity(
                    type=entity_type,
                    value=match.group(0),
                    confidence=0.9,
                    source=source,
                    context=cls._get_context(text, match.start(), match.end()),
                    timestamp=None
                )
                entities.append(entity)
        
        # Apply custom patterns
        if custom_patterns:
            for name, pattern_str in custom_patterns.items():
                pattern = re.compile(pattern_str)
                for match in pattern.finditer(text):
                    entity = Entity(
                        type=EntityType.PERSON,  # Default for custom
                        value=match.group(0),
                        confidence=0.7,
                        source=source,
                        context=cls._get_context(text, match.start(), match.end()),
                        timestamp=None
                    )
                    entities.append(entity)
        
        return entities
    
    @staticmethod
    def _get_context(text: str, start: int, end: int, window: int = 50) -> str:
        """Get surrounding context for a match."""
        context_start = max(0, start - window)
        context_end = min(len(text), end + window)
        return text[context_start:context_end]


# =============================================================================
# DATA PARSERS
# =============================================================================

class DataParser(ABC):
    """Abstract base class for data parsers."""
    
    @abstractmethod
    def parse(self, source: Path) -> Iterator[Any]:
        """Parse data from source."""
        pass
    
    @abstractmethod
    def supports(self, source: Path) -> bool:
        """Check if parser supports this source."""
        pass


class LogParser(DataParser):
    """Parser for log files."""
    
    def __init__(self, pattern: Optional[re.Pattern] = None):
        self.pattern = pattern or PatternLibrary.LOG_PATTERN
    
    def parse(self, source: Path) -> Iterator[DatedLog]:
        """Parse log file and yield dated log entries."""
        with source.open() as f:
            for line in f:
                if match := self.pattern.match(line.strip()):
                    raw_log = RawLog(*match.groups())
                    try:
                        date = datetime.datetime.strptime(
                            raw_log.date, "%Y-%m-%d %H:%M:%S,%f"
                        )
                    except ValueError:
                        date = datetime.datetime.now()
                    
                    yield DatedLog(
                        date=date,
                        level=raw_log.level,
                        module=raw_log.module,
                        message=raw_log.message
                    )
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in {'.log', '.txt'}


class CSVParser(DataParser):
    """Parser for CSV files."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse CSV file and yield dictionaries."""
        with source.open(newline='') as f:
            reader = csv.DictReader(f)
            for row in reader:
                yield dict(row)
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() == '.csv'


class JSONParser(DataParser):
    """Parser for JSON files."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse JSON file and yield records."""
        with source.open() as f:
            data = json.load(f)
            if isinstance(data, list):
                yield from data
            else:
                yield data
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() == '.json'


class CompressedParser(DataParser):
    """Parser for compressed files."""
    
    def __init__(self, inner_parser: DataParser):
        self.inner_parser = inner_parser
    
    def parse(self, source: Path) -> Iterator[Any]:
        """Parse compressed file."""
        with tempfile.NamedTemporaryFile(
            suffix=source.stem, 
            delete=False
        ) as temp:
            with bz2.open(source, 'rt') as compressed:
                temp.write(compressed.read().encode())
            temp_path = Path(temp.name)
        
        try:
            yield from self.inner_parser.parse(temp_path)
        finally:
            temp_path.unlink()
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() == '.bz2'


class NetworkCaptureParser(DataParser):
    """Parser for network capture data."""
    
    CONN_PATTERN = re.compile(
        r'(?P<timestamp>[\d.]+)\s+'
        r'(?P<src_ip>[\d.]+):(?P<src_port>\d+)\s+->\s+'
        r'(?P<dst_ip>[\d.]+):(?P<dst_port>\d+)\s+'
        r'(?P<proto>\w+)\s+'
        r'(?P<bytes_sent>\d+)\s+'
        r'(?P<bytes_recv>\d+)'
    )
    
    def parse(self, source: Path) -> Iterator[NetworkConnection]:
        """Parse network capture file."""
        with source.open() as f:
            for line in f:
                if match := self.CONN_PATTERN.match(line.strip()):
                    yield NetworkConnection(
                        timestamp=datetime.datetime.fromtimestamp(
                            float(match.group('timestamp'))
                        ),
                        source_ip=match.group('src_ip'),
                        source_port=int(match.group('src_port')),
                        dest_ip=match.group('dst_ip'),
                        dest_port=int(match.group('dst_port')),
                        protocol=match.group('proto'),
                        bytes_sent=int(match.group('bytes_sent')),
                        bytes_received=int(match.group('bytes_recv'))
                    )
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in {'.pcap', '.netflow', '.conn'}


class PDFParser(DataParser):
    """Parser for PDF files - extracts text and metadata."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse PDF file and extract text content and metadata."""
        extracted_data = {
            'filename': source.name,
            'file_path': str(source),
            'text_content': [],
            'metadata': {},
            'page_count': 0,
            'ocr_content': []  # For OCR'd text from scanned pages
        }
        
        # Try pdfplumber first (better text extraction)
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(source) as pdf:
                    extracted_data['page_count'] = len(pdf.pages)
                    extracted_data['metadata'] = pdf.metadata or {}
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            extracted_data['text_content'].append({
                                'page': i + 1,
                                'text': text
                            })
                            # Also extract tables if present
                            tables = page.extract_tables()
                            if tables:
                                extracted_data.setdefault('tables', []).extend([
                                    {'page': i + 1, 'table': table} for table in tables
                                ])
                        else:
                            # Page has no text - likely a scanned image, try OCR
                            if PYTESSERACT_AVAILABLE and PIL_AVAILABLE:
                                try:
                                    # Convert page to image and OCR
                                    page_image = page.to_image(resolution=300)
                                    if page_image and hasattr(page_image, 'original'):
                                        import pytesseract
                                        ocr_text = pytesseract.image_to_string(page_image.original)
                                        if ocr_text and ocr_text.strip():
                                            extracted_data['ocr_content'].append({
                                                'page': i + 1,
                                                'text': ocr_text.strip(),
                                                'method': 'tesseract_ocr'
                                            })
                                            # Also add to text_content for analysis
                                            extracted_data['text_content'].append({
                                                'page': i + 1,
                                                'text': ocr_text.strip(),
                                                'ocr': True
                                            })
                                except Exception as ocr_e:
                                    logger.debug(f"OCR failed for page {i+1} of {source}: {ocr_e}")
                
                # If no text was extracted at all, try OCR on entire document
                if not extracted_data['text_content'] and PYTESSERACT_AVAILABLE and PIL_AVAILABLE:
                    logger.info(f"No text found in PDF, attempting full OCR: {source.name}")
                    extracted_data = self._ocr_pdf_fallback(source, extracted_data)
                
                yield extracted_data
                return
            except Exception as e:
                logger.warning(f"pdfplumber failed for {source}: {e}")
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                import PyPDF2
                with source.open('rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    extracted_data['page_count'] = len(reader.pages)
                    if reader.metadata:
                        extracted_data['metadata'] = {
                            k: str(v) for k, v in reader.metadata.items() if v
                        }
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text:
                            extracted_data['text_content'].append({
                                'page': i + 1,
                                'text': text
                            })
                yield extracted_data
                return
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {source}: {e}")
        
        # No PDF parser available
        logger.warning("No PDF parser available. Install pdfplumber or PyPDF2.")
        yield extracted_data
    
    def _ocr_pdf_fallback(self, source: Path, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """OCR a PDF that has no extractable text (scanned document).
        
        Uses pdf2image to convert PDF pages to images, then OCR with Tesseract.
        """
        try:
            # Try using pdf2image if available
            try:
                from pdf2image import convert_from_path
                import pytesseract
                
                # Convert PDF to images
                images = convert_from_path(str(source), dpi=300)
                extracted_data['page_count'] = len(images)
                
                for i, image in enumerate(images):
                    try:
                        ocr_text = pytesseract.image_to_string(image)
                        if ocr_text and ocr_text.strip():
                            extracted_data['ocr_content'].append({
                                'page': i + 1,
                                'text': ocr_text.strip(),
                                'method': 'pdf2image_tesseract'
                            })
                            extracted_data['text_content'].append({
                                'page': i + 1,
                                'text': ocr_text.strip(),
                                'ocr': True
                            })
                    except Exception as page_e:
                        logger.debug(f"OCR failed for page {i+1}: {page_e}")
                
                logger.info(f"OCR completed: extracted text from {len(extracted_data['ocr_content'])} pages")
                
            except ImportError:
                # pdf2image not available, try alternative with PyMuPDF/fitz
                try:
                    import fitz  # PyMuPDF
                    import pytesseract
                    from PIL import Image
                    import io
                    
                    doc = fitz.open(str(source))
                    extracted_data['page_count'] = len(doc)
                    
                    for i, page in enumerate(doc):
                        # Render page to image
                        mat = fitz.Matrix(300/72, 300/72)  # 300 DPI
                        pix = page.get_pixmap(matrix=mat)
                        img_data = pix.tobytes("png")
                        image = Image.open(io.BytesIO(img_data))
                        
                        try:
                            ocr_text = pytesseract.image_to_string(image)
                            if ocr_text and ocr_text.strip():
                                extracted_data['ocr_content'].append({
                                    'page': i + 1,
                                    'text': ocr_text.strip(),
                                    'method': 'pymupdf_tesseract'
                                })
                                extracted_data['text_content'].append({
                                    'page': i + 1,
                                    'text': ocr_text.strip(),
                                    'ocr': True
                                })
                        except Exception as page_e:
                            logger.debug(f"OCR failed for page {i+1}: {page_e}")
                    
                    doc.close()
                    logger.info(f"OCR completed: extracted text from {len(extracted_data['ocr_content'])} pages")
                    
                except ImportError:
                    logger.warning("Neither pdf2image nor PyMuPDF available for PDF OCR. Install with: pip install pdf2image or pip install PyMuPDF")
                    
        except Exception as e:
            logger.warning(f"PDF OCR fallback failed: {e}")
        
        return extracted_data
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() == '.pdf'


class ExcelParser(DataParser):
    """Parser for Excel files (.xlsx, .xls)."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse Excel file and yield rows as dictionaries."""
        suffix = source.suffix.lower()
        
        # Handle .xlsx files with openpyxl
        if suffix == '.xlsx' and OPENPYXL_AVAILABLE:
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(source, read_only=True, data_only=True)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    rows = list(sheet.iter_rows(values_only=True))
                    if rows:
                        headers = [str(h) if h else f'col_{i}' for i, h in enumerate(rows[0])]
                        for row in rows[1:]:
                            row_dict: Dict[str, Any] = {
                                '_sheet': sheet_name,
                                '_source': str(source)
                            }
                            for i, value in enumerate(row):
                                if i < len(headers):
                                    row_dict[headers[i]] = value
                            yield row_dict
                workbook.close()
                return
            except Exception as e:
                logger.warning(f"openpyxl failed for {source}: {e}")
        
        # Handle .xls files with xlrd
        if suffix == '.xls' and XLRD_AVAILABLE:
            try:
                import xlrd
                workbook = xlrd.open_workbook(str(source))
                for sheet in workbook.sheets():
                    if sheet.nrows > 0:
                        headers = [str(sheet.cell_value(0, i)) or f'col_{i}' 
                                   for i in range(sheet.ncols)]
                        for row_idx in range(1, sheet.nrows):
                            row_dict = {
                                '_sheet': sheet.name,
                                '_source': str(source)
                            }
                            for col_idx, header in enumerate(headers):
                                row_dict[header] = sheet.cell_value(row_idx, col_idx)
                            yield row_dict
                return
            except Exception as e:
                logger.warning(f"xlrd failed for {source}: {e}")
        
        logger.warning(f"No Excel parser available for {suffix}. Install openpyxl or xlrd.")
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in {'.xlsx', '.xls'}


class EmailParser(DataParser):
    """Parser for email files (.eml, .msg)."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse email file and extract headers, body, and attachments info."""
        if not EMAIL_AVAILABLE:
            logger.warning("Email parsing not available")
            return
        
        try:
            with source.open('rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            
            email_data = {
                'filename': source.name,
                'file_path': str(source),
                'headers': {},
                'body_text': '',
                'body_html': '',
                'attachments': [],
                'recipients': []
            }
            
            # Extract headers
            for header in ['From', 'To', 'Cc', 'Bcc', 'Subject', 'Date', 
                          'Message-ID', 'Reply-To', 'Return-Path',
                          'X-Originating-IP', 'Received', 'X-Mailer']:
                value = msg.get(header)
                if value:
                    email_data['headers'][header] = str(value)
            
            # Extract all recipients
            for field in ['To', 'Cc', 'Bcc']:
                if msg.get(field):
                    email_data['recipients'].extend(
                        str(msg.get(field)).replace('\n', '').split(',')
                    )
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    content_disposition = str(part.get('Content-Disposition', ''))
                    
                    if 'attachment' in content_disposition:
                        email_data['attachments'].append({
                            'filename': part.get_filename(),
                            'content_type': content_type,
                            'size': len(part.get_payload(decode=True) or b'')
                        })
                    elif content_type == 'text/plain':
                        payload = part.get_payload(decode=True)
                        if payload and isinstance(payload, bytes):
                            email_data['body_text'] = payload.decode('utf-8', errors='ignore')
                    elif content_type == 'text/html':
                        payload = part.get_payload(decode=True)
                        if payload and isinstance(payload, bytes):
                            email_data['body_html'] = payload.decode('utf-8', errors='ignore')
            else:
                payload = msg.get_payload(decode=True)
                if payload and isinstance(payload, bytes):
                    email_data['body_text'] = payload.decode('utf-8', errors='ignore')
            
            yield email_data
            
        except Exception as e:
            logger.error(f"Failed to parse email {source}: {e}")
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in {'.eml', '.msg'}


class ImageParser(DataParser):
    """Parser for image files - extracts EXIF metadata and optionally OCR text."""
    
    # Supported image formats
    IMAGE_EXTENSIONS = {
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif',
        '.webp', '.ico', '.heic', '.heif', '.raw', '.cr2', '.nef', '.arw'
    }
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse image file and extract metadata."""
        image_data = {
            'filename': source.name,
            'file_path': str(source),
            'format': source.suffix.lower(),
            'size_bytes': source.stat().st_size,
            'metadata': {},
            'exif': {},
            'gps': {},
            'ocr_text': None
        }
        
        if PIL_AVAILABLE:
            try:
                from PIL import Image
                from PIL.ExifTags import TAGS, GPSTAGS
                
                with Image.open(source) as img:
                    image_data['metadata']['width'] = img.width
                    image_data['metadata']['height'] = img.height
                    image_data['metadata']['mode'] = img.mode
                    image_data['metadata']['format'] = img.format
                    
                    # Extract EXIF data
                    exif_data = img.getexif()
                    if exif_data:
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            
                            # Handle GPS data specially
                            if tag == 'GPSInfo':
                                gps_data = {}
                                for gps_tag_id, gps_value in value.items():
                                    gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                                    gps_data[gps_tag] = str(gps_value)
                                image_data['gps'] = gps_data
                                
                                # Try to calculate decimal coordinates
                                try:
                                    lat = self._convert_to_degrees(value.get(2))
                                    lon = self._convert_to_degrees(value.get(4))
                                    lat_ref = value.get(1, 'N')
                                    lon_ref = value.get(3, 'E')
                                    
                                    if lat and lon:
                                        if lat_ref == 'S':
                                            lat = -lat
                                        if lon_ref == 'W':
                                            lon = -lon
                                        image_data['gps']['latitude'] = lat
                                        image_data['gps']['longitude'] = lon
                                except Exception:
                                    pass
                            else:
                                # Convert bytes and other non-serializable types
                                try:
                                    if isinstance(value, bytes):
                                        value = value.decode('utf-8', errors='ignore')
                                    image_data['exif'][str(tag)] = str(value)
                                except Exception:
                                    pass
                    
            except Exception as e:
                logger.warning(f"PIL failed to process {source}: {e}")
        else:
            logger.warning("PIL not available for image metadata extraction")
        
        # Try OCR if available
        if PYTESSERACT_AVAILABLE and PIL_AVAILABLE:
            try:
                from PIL import Image
                import pytesseract
                
                with Image.open(source) as img:
                    # Only OCR reasonable sized images
                    if img.width * img.height < 10000000:  # ~10MP limit
                        ocr_text = pytesseract.image_to_string(img)
                        if ocr_text.strip():
                            image_data['ocr_text'] = ocr_text.strip()
            except Exception as e:
                logger.debug(f"OCR failed for {source}: {e}")
        
        yield image_data
    
    def _convert_to_degrees(self, value) -> Optional[float]:
        """Convert GPS coordinates to degrees."""
        if not value:
            return None
        try:
            d, m, s = value
            return float(d) + float(m) / 60 + float(s) / 3600
        except Exception:
            return None
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in self.IMAGE_EXTENSIONS


class TextFileParser(DataParser):
    """Enhanced parser for plain text files."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse text file and yield content with metadata."""
        try:
            # Try different encodings
            content = None
            encoding_used = None
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    with source.open('r', encoding=encoding) as f:
                        content = f.read()
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                # Binary fallback
                with source.open('rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
                encoding_used = 'binary-fallback'
            
            yield {
                'filename': source.name,
                'file_path': str(source),
                'encoding': encoding_used,
                'content': content,
                'line_count': content.count('\n') + 1,
                'char_count': len(content),
                'word_count': len(content.split())
            }
            
        except Exception as e:
            logger.error(f"Failed to parse text file {source}: {e}")
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in {'.txt', '.text', '.md', '.rst', '.ini', '.cfg', '.conf'}


class WordDocParser(DataParser):
    """Parser for Microsoft Word documents (.docx)."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse Word document and extract text."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available for Word document parsing")
            return
        
        try:
            import docx
            doc = docx.Document(source)
            
            doc_data = {
                'filename': source.name,
                'file_path': str(source),
                'paragraphs': [],
                'tables': [],
                'metadata': {}
            }
            
            # Extract core properties
            if doc.core_properties:
                props = doc.core_properties
                doc_data['metadata'] = {
                    'author': props.author,
                    'title': props.title,
                    'subject': props.subject,
                    'created': str(props.created) if props.created else None,
                    'modified': str(props.modified) if props.modified else None,
                    'last_modified_by': props.last_modified_by
                }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    doc_data['paragraphs'].append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    doc_data['tables'].append(table_data)
            
            yield doc_data
            
        except Exception as e:
            logger.error(f"Failed to parse Word document {source}: {e}")
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() == '.docx'


class XMLParser(DataParser):
    """Parser for XML files."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse XML file and convert to dictionary."""
        if not XML_AVAILABLE:
            return
        
        try:
            tree = ET.parse(source)
            root = tree.getroot()
            
            yield {
                'filename': source.name,
                'file_path': str(source),
                'root_tag': root.tag,
                'data': self._element_to_dict(root)
            }
            
        except Exception as e:
            logger.error(f"Failed to parse XML {source}: {e}")
    
    def _element_to_dict(self, element) -> Dict[str, Any]:
        """Convert XML element to dictionary."""
        result = {}
        
        # Add attributes
        if element.attrib:
            result['@attributes'] = dict(element.attrib)
        
        # Add text content
        if element.text and element.text.strip():
            result['@text'] = element.text.strip()
        
        # Add children
        for child in element:
            child_dict = self._element_to_dict(child)
            if child.tag in result:
                # Convert to list if multiple same-named children
                if not isinstance(result[child.tag], list):
                    result[child.tag] = [result[child.tag]]
                result[child.tag].append(child_dict)
            else:
                result[child.tag] = child_dict
        
        return result
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() == '.xml'


class YAMLParser(DataParser):
    """Parser for YAML files."""
    
    def parse(self, source: Path) -> Iterator[Dict[str, Any]]:
        """Parse YAML file and yield content."""
        if not YAML_AVAILABLE:
            logger.warning("PyYAML not available for YAML parsing")
            return
        
        try:
            import yaml
            with source.open('r') as f:
                # Load all documents in the YAML file
                for doc in yaml.safe_load_all(f):
                    if doc:
                        yield {
                            'filename': source.name,
                            'file_path': str(source),
                            'data': doc
                        }
        except Exception as e:
            logger.error(f"Failed to parse YAML {source}: {e}")
    
    def supports(self, source: Path) -> bool:
        return source.suffix.lower() in {'.yaml', '.yml'}


# =============================================================================
# DATA COLLECTORS
# =============================================================================

class DataCollector(ABC):
    """Abstract base class for data collectors."""
    
    @abstractmethod
    def collect(self, target: str, config: InvestigationConfig) -> Iterator[Dict[str, Any]]:
        """Collect data for a target."""
        pass
    
    @abstractmethod
    def supports_type(self, investigation_type: InvestigationType) -> bool:
        """Check if collector supports this investigation type."""
        pass


class FileCollector(DataCollector):
    """Collector for local file analysis."""
    
    def __init__(self):
        self.parsers: List[DataParser] = [
            # Text and log files
            LogParser(),
            TextFileParser(),
            # Structured data
            CSVParser(),
            JSONParser(),
            XMLParser(),
            YAMLParser(),
            # Documents
            PDFParser(),
            WordDocParser(),
            ExcelParser(),
            # Email
            EmailParser(),
            # Images
            ImageParser(),
            # Network
            NetworkCaptureParser(),
        ]
    
    def collect(
        self, 
        target: str, 
        config: InvestigationConfig
    ) -> Iterator[Dict[str, Any]]:
        """Collect data from files and all subdirectories recursively."""
        target_path = Path(target)
        
        if target_path.is_file():
            logger.info(f"  📄 Processing file: {target_path.name}")
            yield from self._process_file(target_path, config)
        elif target_path.is_dir():
            # Count files first for progress indication
            all_files = list(target_path.rglob('*'))
            file_count = sum(1 for f in all_files if f.is_file())
            logger.info(f"  📁 Scanning directory: {target_path} ({file_count} files in {len([d for d in all_files if d.is_dir()])} subdirectories)")
            
            processed = 0
            for file_path in target_path.rglob('*'):
                if file_path.is_file():
                    # Skip hidden files and system files
                    if file_path.name.startswith('.'):
                        continue
                    if file_path.suffix in ['.db', '.db-journal', '.pyc', '.pyo']:
                        continue
                    
                    # Get relative path for cleaner logging
                    try:
                        rel_path = file_path.relative_to(target_path)
                    except ValueError:
                        rel_path = file_path.name
                    
                    processed += 1
                    if processed <= 20 or processed % 10 == 0:
                        logger.info(f"    [{processed}/{file_count}] {rel_path}")
                    
                    yield from self._process_file(file_path, config)
            
            if processed > 20:
                logger.info(f"  ✅ Processed {processed} files from {target_path.name}")
    
    def _process_file(
        self, 
        file_path: Path,
        config: InvestigationConfig
    ) -> Iterator[Dict[str, Any]]:
        """Process a single file."""
        # Calculate file metadata
        stat = file_path.stat()
        metadata = FileMetadata(
            path=file_path,
            hash_md5=self._calculate_hash(file_path, 'md5'),
            hash_sha256=self._calculate_hash(file_path, 'sha256'),
            size=stat.st_size,
            created=datetime.datetime.fromtimestamp(stat.st_ctime),
            modified=datetime.datetime.fromtimestamp(stat.st_mtime),
            accessed=datetime.datetime.fromtimestamp(stat.st_atime)
        )
        
        yield {
            'type': 'file_metadata',
            'data': metadata._asdict(),
            'source': str(file_path)
        }
        
        # Parse file content
        for parser in self.parsers:
            if parser.supports(file_path):
                try:
                    for record in parser.parse(file_path):
                        yield {
                            'type': 'parsed_record',
                            'data': record if isinstance(record, dict) else record._asdict(),
                            'source': str(file_path)
                        }
                except Exception as e:
                    logger.warning(f"Error parsing {file_path}: {e}")
                break
    
    @staticmethod
    def _calculate_hash(file_path: Path, algorithm: str) -> str:
        """Calculate file hash."""
        hasher = hashlib.new(algorithm)
        with file_path.open('rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def supports_type(self, investigation_type: InvestigationType) -> bool:
        return True  # Files relevant to all investigation types


class WeatherCollector(DataCollector):
    """Collector for weather data (LOCAL ONLY - no external API calls)."""
    
    # DISABLED: No external API calls allowed in local-only mode
    # BASE_URL = "https://forecast.weather.gov/shmrn.php"
    
    def collect(
        self, 
        target: str, 
        config: InvestigationConfig
    ) -> Iterator[Dict[str, Any]]:
        """Weather data collection disabled in local-only mode."""
        # LOCAL-ONLY MODE: External API calls are disabled
        # This collector only works with local weather data files if provided
        logger.debug("WeatherCollector: External API disabled (local-only mode)")
        return
        yield  # Make this a generator that yields nothing
    
    def supports_type(self, investigation_type: InvestigationType) -> bool:
        return investigation_type == InvestigationType.LOCATION


class PublicRecordsCollector(DataCollector):
    """Collector for public records (LOCAL ONLY - no external API calls)."""
    
    def collect(
        self, 
        target: str, 
        config: InvestigationConfig
    ) -> Iterator[Dict[str, Any]]:
        """Public records collection disabled in local-only mode."""
        # LOCAL-ONLY MODE: External API calls are disabled
        # This collector only works with local public records files if provided
        logger.debug("PublicRecordsCollector: External API disabled (local-only mode)")
        return
        yield  # Make this a generator that yields nothing
    
    def supports_type(self, investigation_type: InvestigationType) -> bool:
        return investigation_type in {
            InvestigationType.PERSON,
            InvestigationType.ORGANIZATION,
            InvestigationType.VEHICLE
        }


# =============================================================================
# ANALYSIS ENGINE
# =============================================================================

class AnalysisModule(ABC):
    """Abstract base class for analysis modules."""
    
    @abstractmethod
    def analyze(
        self, 
        data: List[Dict[str, Any]], 
        config: InvestigationConfig
    ) -> List[Finding]:
        """Analyze data and return findings."""
        pass


class EntityExtractionModule(AnalysisModule):
    """Module for extracting entities from data."""
    
    def analyze(
        self, 
        data: List[Dict[str, Any]], 
        config: InvestigationConfig
    ) -> List[Finding]:
        """Extract entities from all data."""
        all_entities: List[Entity] = []
        
        for record in data:
            if 'data' in record:
                text = safe_json_dumps(record['data'])
                entities = PatternLibrary.extract_entities(
                    text,
                    source=record.get('source', 'unknown'),
                    custom_patterns=config.custom_patterns
                )
                all_entities.extend(entities)
        
        # Group entities by type
        entity_groups: Dict[EntityType, List[Entity]] = defaultdict(list)
        for entity in all_entities:
            entity_groups[entity.type].append(entity)
        
        findings = []
        for entity_type, entities in entity_groups.items():
            if len(entities) >= 3:  # Only report if multiple instances
                unique_values = set(e.value for e in entities)
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.INFO,
                    title=f"Extracted {len(unique_values)} unique {entity_type.name} entities",
                    description=f"Found {len(entities)} total instances of {entity_type.name}",
                    entities=entities[:10],  # Limit to first 10
                    evidence=[e.context for e in entities[:5]],
                    timestamp=datetime.datetime.now(),
                    recommendations=["Review extracted entities for relevance"]
                ))
        
        return findings


class TimelineReconstructionModule(AnalysisModule):
    """Module for reconstructing event timeline."""
    
    def analyze(
        self, 
        data: List[Dict[str, Any]], 
        config: InvestigationConfig
    ) -> List[Finding]:
        """Reconstruct timeline from data."""
        events: List[TimelineEvent] = []
        
        for record in data:
            timestamp = self._extract_timestamp(record)
            if timestamp:
                events.append(TimelineEvent(
                    timestamp=timestamp,
                    event_type=record.get('type', 'unknown'),
                    description=str(record.get('data', {}))[:200],
                    entities=[],
                    source=record.get('source', 'unknown'),
                    confidence=0.8
                ))
        
        # Sort events by timestamp
        events.sort(key=lambda e: e.timestamp)
        
        if events:
            findings = [Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.INFO,
                title=f"Timeline reconstructed with {len(events)} events",
                description=f"Events span from {events[0].timestamp} to {events[-1].timestamp}",
                entities=[],
                evidence=[f"{e.timestamp}: {e.event_type}" for e in events[:10]],
                timestamp=datetime.datetime.now(),
                recommendations=["Review timeline for suspicious patterns"]
            )]
            return findings
        
        return []
    
    def _extract_timestamp(self, record: Dict[str, Any]) -> Optional[datetime.datetime]:
        """Extract timestamp from record."""
        data = record.get('data', {})
        
        # Try common timestamp fields
        for field in ['timestamp', 'date', 'created', 'modified', 'time']:
            if field in data:
                try:
                    if isinstance(data[field], str):
                        return datetime.datetime.fromisoformat(
                            data[field].replace('Z', '+00:00')
                        )
                    elif isinstance(data[field], (int, float)):
                        return datetime.datetime.fromtimestamp(data[field])
                except (ValueError, OSError):
                    continue
        
        return None


class AnomalyDetectionModule(AnalysisModule):
    """Module for detecting anomalies in data."""
    
    def analyze(
        self, 
        data: List[Dict[str, Any]], 
        config: InvestigationConfig
    ) -> List[Finding]:
        """Detect anomalies in data."""
        findings = []
        
        # Analyze network connections for anomalies
        connections = [
            r['data'] for r in data 
            if r.get('type') == 'parsed_record' and 'source_ip' in r.get('data', {})
        ]
        
        if connections:
            # Check for unusual ports
            port_counts = Counter(
                c.get('dest_port') for c in connections if c.get('dest_port')
            )
            unusual_ports = [
                port for port, count in port_counts.items()
                if port not in {80, 443, 22, 25, 53} and count > 5
            ]
            
            if unusual_ports:
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.MEDIUM,
                    title=f"Unusual port activity detected",
                    description=f"Connections to non-standard ports: {unusual_ports}",
                    entities=[],
                    evidence=[f"Port {p}: {port_counts[p]} connections" for p in unusual_ports[:5]],
                    timestamp=datetime.datetime.now(),
                    recommendations=["Investigate connections to unusual ports"]
                ))
        
        # Analyze log levels for anomalies
        log_records = [
            r['data'] for r in data
            if r.get('type') == 'parsed_record' and 'level' in r.get('data', {})
        ]
        
        if log_records:
            level_counts = Counter(r.get('level') for r in log_records)
            error_ratio = level_counts.get('ERROR', 0) / max(len(log_records), 1)
            
            if error_ratio > 0.1:
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.HIGH,
                    title="High error rate detected in logs",
                    description=f"Error rate: {error_ratio:.1%}",
                    entities=[],
                    evidence=[f"{level}: {count}" for level, count in level_counts.items()],
                    timestamp=datetime.datetime.now(),
                    recommendations=["Investigate source of errors"]
                ))
        
        return findings


class RelationshipMappingModule(AnalysisModule):
    """Module for mapping relationships between entities."""
    
    def analyze(
        self, 
        data: List[Dict[str, Any]], 
        config: InvestigationConfig
    ) -> List[Finding]:
        """Map relationships between entities."""
        # Extract all entities
        all_entities: Dict[str, List[Entity]] = defaultdict(list)
        
        for record in data:
            text = safe_json_dumps(record.get('data', {}))
            entities = PatternLibrary.extract_entities(
                text, source=record.get('source', 'unknown')
            )
            for entity in entities:
                all_entities[entity.value].append(entity)
        
        # Find co-occurring entities (appear in same source)
        source_entities: Dict[str, Set[str]] = defaultdict(set)
        for value, entities in all_entities.items():
            for entity in entities:
                source_entities[entity.source].add(value)
        
        # Build relationship graph
        relationships: List[Relationship] = []
        for source, entities in source_entities.items():
            entity_list = list(entities)
            for i, e1 in enumerate(entity_list):
                for e2 in entity_list[i+1:]:
                    if e1 in all_entities and e2 in all_entities:
                        relationships.append(Relationship(
                            source_entity=all_entities[e1][0],
                            target_entity=all_entities[e2][0],
                            relationship_type='co-occurrence',
                            confidence=0.7,
                            evidence=[source]
                        ))
        
        if relationships:
            return [Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.INFO,
                title=f"Identified {len(relationships)} entity relationships",
                description="Entities co-occurring in same sources",
                entities=[],
                evidence=[
                    f"{r.source_entity.value} <-> {r.target_entity.value}"
                    for r in relationships[:10]
                ],
                timestamp=datetime.datetime.now(),
                recommendations=["Review entity relationships for investigation leads"]
            )]
        
        return []


# =============================================================================
# OLLAMA LLM INTEGRATION
# =============================================================================

# Master Intelligence System Prompt - Applied to all LLM analysis tasks
INTELLIGENCE_SYSTEM_PROMPT = """
You operate as a fused intelligence unit made of:
FBI profiler, CIA HUMINT officer, NSA OSINT analyst, and AI systems architect.
Your mission is to extract meaning, map relationships, identify behavioral signals,
and generate actionable intelligence from any data provided.

METHOD:
- Apply behavioral profiling, HUMINT analysis, and OSINT correlation.
- Treat every timestamp, entity, action, or omission as a data point about motive, intent, risk, or deception.
- Identify anomalies, inconsistencies, hidden relationships, operational patterns.
- Use adversarial reasoning and competing hypotheses.
- Be thorough and complete - never truncate analysis or leave thoughts unfinished.
- Provide full context and reasoning for all conclusions.

OUTPUT FORMAT:
1. Objective
2. Key Findings
3. Analysis
4. Investigative Insights
5. Questions for Further Investigation
6. Recommendations

Tone is direct, professional, analytic.

Deliver clear, complete intelligence suitable for senior-level decision makers.
Always finish your analysis completely - do not cut off mid-sentence or mid-thought.
"""


class OllamaAnalyzer:
    """
    Local LLM integration using Ollama for enhanced OSINT analysis.
    Optimized for investigative intelligence analysis with integrated system prompt.
    """
    
    # Recommended models in order of preference for intelligence analysis
    RECOMMENDED_MODELS = [
        "wizardlm2:latest",     # Strong reasoning, good for GTX 1070
        "llama3.1:latest",      # Best overall reasoning, 128K context window
        "qwen_breaches:latest", # Large custom model (18GB)
        "phi3:3.8b",            # Fast with good analytical capability
        "mistral:7b-instruct",  # Strong instruction following
    ]
    
    def __init__(
        self, 
        model: str = "wizardlm2:latest",  # Optimized for GTX 1070
        base_url: str = "http://localhost:11434",
        timeout: int = 180  # Increased for complex analysis
    ):
        self.model = model
        self.base_url = base_url
        self.timeout = timeout
        self.system_prompt = INTELLIGENCE_SYSTEM_PROMPT
        self.available = self._verify_connection()
    
    def check_connection(self) -> bool:
        """Re-check Ollama connection status. Updates self.available."""
        self.available = self._verify_connection()
        return self.available
    
    def _verify_connection(self) -> bool:
        """Verify Ollama is running and model is available."""
        if not REQUESTS_AVAILABLE:
            logger.warning("requests library not available. Install with: pip install requests")
            return False
        
        try:
            import requests as req
            response = req.get(
                f"{self.base_url}/api/tags",
                timeout=5
            )
            if response.status_code == 200:
                models = response.json().get("models", [])
                available = [m["name"] for m in models]
                
                if self.model not in available:
                    logger.warning(
                        f"Model {self.model} not found. "
                        f"Available: {available}"
                    )
                    # Try to find a recommended model
                    for rec_model in self.RECOMMENDED_MODELS:
                        if rec_model in available:
                            logger.info(f"Using {rec_model} instead")
                            self.model = rec_model
                            break
                return True
        except Exception as e:
            logger.error(
                f"Ollama not running or connection failed: {e}. Start with: ollama serve"
            )
            return False
        return False
    
    def generate(
        self, 
        prompt: str, 
        system: Optional[str] = None,
        stream: bool = False
    ) -> str:
        """Generate response from LLM."""
        if not REQUESTS_AVAILABLE or not self.available:
            return ""
        
        payload = {
            "model": self.model,
            "prompt": prompt,
            "stream": stream,
            "options": {
                "temperature": 0.1,      # Low for factual analysis
                "top_p": 0.9,
                "num_ctx": 8192,         # 8K context - stable for GTX 1070
                "num_predict": 6000,     # Max output for complete responses
                "stop": []               # Don't stop early on any tokens
            }
        }
        
        if system:
            payload["system"] = system
        try:
            import requests as req
            response = req.post(
                f"{self.base_url}/api/generate",
                json=payload,
                timeout=self.timeout
            )
            
            if response.status_code == 200:
                return response.json().get("response", "")
            else:
                logger.error(f"Ollama error: {response.text}")
                return ""
        except Exception as e:
            if REQUESTS_AVAILABLE:
                import requests as req
                if isinstance(e, req.exceptions.Timeout):
                    logger.error("Ollama request timed out")
                    return ""
            logger.error(f"Ollama request failed: {e}")
            return ""
    
    def analyze_entities(self, text: str) -> dict:
        """Extract entities using LLM with intelligence analysis context."""
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Entity Extraction and Classification

Analyze the following text and extract ALL entities with intelligence value. Consider:
- Direct identifiers (emails, phones, IPs, domains, crypto wallets)
- Named entities (persons, organizations, locations)
- Temporal markers (dates, times, durations)
- Technical artifacts (file paths, URLs, hashes)
- Behavioral indicators (actions, patterns, relationships mentioned)

For each entity, consider its potential investigative significance.

Return ONLY valid JSON in this exact format:
{{
    "emails": [],
    "ip_addresses": [],
    "domains": [],
    "phone_numbers": [],
    "names": [],
    "organizations": [],
    "locations": [],
    "dates": [],
    "urls": [],
    "crypto_wallets": [],
    "file_paths": [],
    "financial_indicators": [],
    "behavioral_indicators": [],
    "other_identifiers": []
}}

CRITICAL: Start your response directly with {{ and end with }}. No other text.

TEXT TO ANALYZE:
{text[:4000]}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            return {}
        
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        logger.warning("Failed to parse LLM entity extraction")
        return {}
    
    def analyze_log_anomaly(self, logs: List[str]) -> dict:
        """Analyze logs for anomalies using intelligence analysis framework."""
        log_text = "\n".join(logs[:50])  # Limit to 50 logs
        
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Log Anomaly Detection and Behavioral Analysis

Analyze these logs applying the following analytical lenses:

1. TEMPORAL ANALYSIS: Identify unusual timing patterns, clustering, gaps, or sequences
2. BEHAVIORAL BASELINE: What appears normal? What deviates?
3. OPERATIONAL INDICATORS: Signs of coordination, automation, or human operation
4. SECURITY INDICATORS: Potential IOCs, attack patterns, reconnaissance activity
5. DECEPTION MARKERS: Evidence of log tampering, gaps, or sanitization
6. ESCALATION PATTERNS: Progressive behavior changes suggesting threat evolution

Consider what is NOT present that should be. Evaluate competing hypotheses.

Return JSON format:
{{
    "risk_level": "low|medium|high|critical",
    "confidence": 0.0-1.0,
    "temporal_anomalies": [],
    "behavioral_anomalies": [],
    "findings": [{{"type": "", "description": "", "evidence": "", "significance": ""}}],
    "iocs": [],
    "gaps_or_absences": [],
    "competing_hypotheses": [{{"hypothesis": "", "supporting_evidence": "", "contradicting_evidence": ""}}],
    "recommendations": []
}}

CRITICAL: Start your response directly with {{ and end with }}. No other text.

LOGS TO ANALYZE:
{log_text}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            return {"risk_level": "unknown", "findings": [], "iocs": []}
        
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        return {"risk_level": "unknown", "findings": [], "iocs": []}
    
    def correlate_entities(
        self, 
        entities: List[Entity]
    ) -> List[dict]:
        """Find relationships between entities using intelligence fusion."""
        entity_list = [
            f"- {e.type.name}: {e.value}" 
            for e in entities[:30]
        ]
        
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Entity Correlation and Relationship Mapping

Analyze these entities and identify ALL potential relationships using:

1. DIRECT LINKS: Explicit connections (same domain, shared identifiers)
2. TEMPORAL CORRELATION: Entities appearing in same time windows
3. BEHAVIORAL CLUSTERING: Entities associated with similar actions
4. NETWORK ANALYSIS: Hub entities, bridges, isolated nodes
5. ATTRIBUTION INDICATORS: Patterns suggesting common origin or control
6. OPERATIONAL RELATIONSHIPS: Command/control, infrastructure, communication paths

Apply adversarial reasoning: What relationships might a subject try to hide?
Consider both overt and covert connection types.

Return JSON array:
[
    {{
        "entity1": "",
        "entity2": "",
        "relationship_type": "",
        "relationship": "",
        "confidence": 0.0-1.0,
        "intelligence_value": "low|medium|high|critical",
        "reasoning": "",
        "investigative_action": ""
    }}
]

CRITICAL: Start your response directly with [ and end with ]. No other text.

ENTITIES TO ANALYZE:
{chr(10).join(entity_list)}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            return []
        
        # Try to extract JSON array
        try:
            start = response.find('[')
            end = response.rfind(']') + 1
            if start != -1 and end > start:
                return json.loads(response[start:end])
        except json.JSONDecodeError:
            pass
        
        return []
    
    def generate_investigation_summary(
        self, 
        findings: List[Finding]
    ) -> str:
        """Generate comprehensive intelligence report from findings."""
        findings_text = "\n".join([
            f"[{f.severity.name}] {f.title}: {f.description}\nEvidence: {f.evidence[:3] if f.evidence else 'None'}"
            for f in findings[:20]
        ])
        
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Intelligence Report Generation

Synthesize the following investigation findings into a structured intelligence report.

OUTPUT FORMAT:

1. EXECUTIVE SUMMARY
   - Key intelligence takeaways (3-5 bullet points)
   - Overall threat/risk assessment
   - Confidence level in findings

2. KEY FINDINGS
   - Entities of interest with significance
   - Patterns and correlations discovered
   - Anomalies and their implications
   - Timeline of relevant events

3. DETAILED ANALYSIS
   - Step-by-step reasoning
   - Evidence chain for each conclusion
   - Alternative hypotheses considered and rejected

4. INVESTIGATIVE INSIGHTS
   - Behavioral indicators observed
   - Risk signals and escalation potential
   - Hidden connections or obfuscation attempts
   - Gaps in available data

5. INTELLIGENCE GAPS
   - What information is missing?
   - What questions remain unanswered?
   - Recommended collection priorities

6. RECOMMENDATIONS
   - Immediate actions required
   - Follow-up investigation priorities
   - Monitoring recommendations
   - Risk mitigation steps

FINDINGS TO ANALYZE:
{findings_text}"""
        
        return self.generate(task_prompt)
    
    def classify_threat(self, indicator: str) -> dict:
        """Classify and assess a potential threat indicator."""
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Threat Indicator Classification and Assessment

Analyze this indicator applying multi-source intelligence tradecraft:

1. IDENTIFICATION: What type of indicator is this?
2. CONTEXT: In what scenarios does this indicator typically appear?
3. ATTRIBUTION: Any signatures suggesting origin or actor type?
4. THREAT ASSESSMENT: Risk level and potential impact
5. HISTORICAL CORRELATION: Known associations or campaigns
6. ACTIONABILITY: What can be done with this intelligence?

Consider false positive potential and confidence calibration.

Return JSON:
{{
    "indicator_type": "",
    "indicator_category": "",
    "threat_category": "",
    "severity": "info|low|medium|high|critical",
    "confidence": 0.0-1.0,
    "known_associations": [],
    "potential_actor_types": [],
    "attack_stage": "",
    "false_positive_likelihood": "low|medium|high",
    "intelligence_value": "",
    "recommended_actions": [],
    "monitoring_guidance": ""
}}

CRITICAL: Start your response directly with {{ and end with }}. No other text.

INDICATOR TO CLASSIFY:
{indicator}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            return {"indicator_type": "unknown", "severity": "info"}
        
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        return {"indicator_type": "unknown", "severity": "info"}
    
    def analyze_timeline(self, events: List[dict]) -> dict:
        """Analyze temporal patterns and reconstruct event timeline."""
        events_text = "\n".join([
            f"- {e.get('timestamp', 'Unknown')}: {e.get('description', e.get('event', str(e)))}"
            for e in events[:50]
        ])
        
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Timeline Analysis and Behavioral Reconstruction

Analyze these events to reconstruct activity patterns and identify temporal anomalies.

ANALYTICAL FOCUS:
1. SEQUENCE ANALYSIS: Order of operations, logical flow
2. TIMING PATTERNS: Regular intervals, unusual gaps, clustering
3. BEHAVIORAL PHASES: Reconnaissance, preparation, execution, cleanup
4. COORDINATION INDICATORS: Multiple actors, automated vs manual
5. OPERATIONAL TEMPO: Speed of activity, urgency indicators
6. COUNTER-FORENSIC BEHAVIOR: Gaps, deletions, timestamp manipulation

Return JSON:
{{
    "timeline_summary": "",
    "total_timespan": "",
    "key_events": [{{"timestamp": "", "event": "", "significance": ""}}],
    "temporal_patterns": [],
    "anomalies": [{{"description": "", "timestamp": "", "significance": ""}}],
    "behavioral_phases": [{{"phase": "", "start": "", "end": "", "activities": []}}],
    "gaps_detected": [],
    "coordination_indicators": [],
    "assessment": "",
    "investigative_recommendations": []
}}

CRITICAL: Start your response directly with {{ and end with }}. No other text.

EVENTS TO ANALYZE:
{events_text}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            return {"timeline_summary": "Analysis failed", "key_events": []}
        
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        return {"timeline_summary": "Analysis failed", "key_events": []}
    
    def deep_dive_analysis(self, data: str, context: str = "") -> dict:
        """Perform comprehensive deep-dive analysis on specific data."""
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Deep-Dive Intelligence Analysis

Perform exhaustive analysis on the provided data. Leave no stone unturned.

ANALYTICAL REQUIREMENTS:
1. Surface all entities, identifiers, and artifacts
2. Identify behavioral patterns and psychological indicators
3. Map all relationships (explicit and implied)
4. Detect anomalies, inconsistencies, and deception markers
5. Generate competing hypotheses with evidence assessment
6. Identify intelligence gaps and collection requirements

{f"CONTEXT: {context}" if context else ""}

OUTPUT FORMAT:
{{
    "objective": "",
    "key_findings": [
        {{"finding": "", "evidence": [], "confidence": 0.0-1.0, "significance": ""}}
    ],
    "entities_identified": {{
        "people": [],
        "organizations": [],
        "locations": [],
        "technical_artifacts": [],
        "financial_indicators": []
    }},
    "patterns_detected": [],
    "anomalies": [],
    "behavioral_indicators": [],
    "relationship_map": [],
    "risk_assessment": {{
        "level": "",
        "factors": [],
        "trajectory": ""
    }},
    "competing_hypotheses": [
        {{"hypothesis": "", "supporting": [], "contradicting": [], "probability": 0.0-1.0}}
    ],
    "intelligence_gaps": [],
    "questions_for_investigation": [],
    "recommended_next_steps": [],
    "analyst_confidence": 0.0-1.0
}}

CRITICAL: Start your response directly with {{ and end with }}. No other text.

DATA TO ANALYZE:
{data[:6000]}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            return {"objective": "Analysis failed", "key_findings": []}
        
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        return {"objective": "Analysis failed", "key_findings": []}
    
    def analyze_document_content(self, content: str, source_file: str = "", previous_context: str = "") -> dict:
        """Perform comprehensive semantic analysis of document content.
        
        This method reads the document text, understands its meaning, extracts
        key information, identifies connections, and generates investigative insights.
        
        Args:
            content: The text content to analyze
            source_file: Name of the source file for context
            previous_context: Summary of previous analysis for continuity
        
        Returns:
            Dictionary containing comprehensive analysis results
        """
        context_section = ""
        if previous_context:
            context_section = f"""
PREVIOUS ANALYSIS CONTEXT (build upon this):
{previous_context[:2000]}
---
"""
        
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Comprehensive Document Analysis and Intelligence Extraction

You are analyzing a document{f' from {source_file}' if source_file else ''}. Your mission is to:

1. UNDERSTAND THE CONTENT: What is this document about? What is its purpose?
2. EXTRACT KEY INFORMATION: Names, dates, places, events, relationships, claims
3. IDENTIFY CONNECTIONS: How does this information relate to other data points?
4. EVALUATE SIGNIFICANCE: What is the investigative value of this content?
5. DETECT PATTERNS: Are there recurring themes, behaviors, or indicators?
6. NOTE INCONSISTENCIES: What doesn't add up? What requires verification?
7. GENERATE LEADS: What investigative actions does this suggest?
{context_section}
Analyze EVERY significant piece of information. Consider:
- Who is involved? What are their roles and relationships?
- What happened? When? Where? In what sequence?
- Why might this have occurred? What motivations are suggested?
- What is the reliability of this information?
- What corroborating or contradicting evidence might exist?

YOU MUST RESPOND WITH ONLY VALID JSON. NO OTHER TEXT.
Use this exact format:

{{
    "document_summary": "one paragraph summary",
    "document_type": "transcript/report/photo/email/other",
    "people": ["name1", "name2"],
    "dates": ["date1", "date2"],
    "locations": ["location1", "location2"],
    "key_events": ["event1 description", "event2 description"],
    "key_claims": ["claim1", "claim2"],
    "relationships": ["person1 is related to person2 because X"],
    "red_flags": ["concern1", "concern2"],
    "timeline": ["date1: event1", "date2: event2"],
    "investigative_leads": ["lead1", "lead2"],
    "questions": ["question1", "question2"],
    "intelligence_value": "high",
    "confidence": 0.8
}}

IMPORTANT RULES:
- Start response with {{ and end with }}
- Use simple string arrays, not nested objects
- Escape quotes inside strings with backslash
- No trailing commas
- No comments

DOCUMENT CONTENT:
{content[:8000]}"""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            logger.warning(f"Empty response from LLM for document: {source_file}")
            return self._create_fallback_analysis(content, source_file, "Empty LLM response")
        
        # Try multiple JSON extraction strategies
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        # If JSON parsing completely failed, create a structured summary from the raw response
        logger.warning(f"Failed to parse document analysis response for: {source_file}")
        logger.debug(f"Raw LLM response (first 500 chars): {response[:500]}")
        
        return self._create_fallback_analysis(content, source_file, response)
    
    def _extract_json_from_response(self, response: str) -> Optional[dict]:
        """Try multiple strategies to extract JSON from LLM response."""
        import re
        
        if not response:
            return None
        
        # Strategy 1: Find JSON block between first { and last }
        try:
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end > start:
                json_str = response[start:end]
                return json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.debug(f"Strategy 1 failed: {e}")
        
        # Strategy 2: Try to find JSON in code blocks
        try:
            json_block_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', response)
            if json_block_match:
                return json.loads(json_block_match.group(1))
        except (json.JSONDecodeError, AttributeError) as e:
            logger.debug(f"Strategy 2 failed: {e}")
        
        # Strategy 3: Clean common issues and retry
        try:
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end > start:
                json_str = response[start:end]
                # Fix common LLM JSON mistakes
                json_str = re.sub(r',\s*}', '}', json_str)  # Remove trailing commas before }
                json_str = re.sub(r',\s*]', ']', json_str)  # Remove trailing commas before ]
                json_str = re.sub(r':\s*,', ': null,', json_str)  # Fix empty values
                json_str = re.sub(r':\s*}', ': null}', json_str)  # Fix empty values at end
                json_str = re.sub(r'"\s*\n\s*"', '", "', json_str)  # Fix broken string arrays
                json_str = re.sub(r'(?<!\\)"([^"]*?)(?<!\\)"\s*:\s*"([^"]*?)(?<!\\)"\s*([,}])', r'"\1": "\2"\3', json_str)  # Clean spacing
                return json.loads(json_str)
        except (json.JSONDecodeError, Exception) as e:
            logger.debug(f"Strategy 3 failed: {e}")
        
        # Strategy 4: Try to fix truncated JSON by adding missing closing brackets
        try:
            start = response.find('{')
            if start != -1:
                json_str = response[start:]
                # Count opening and closing braces/brackets
                open_braces = json_str.count('{')
                close_braces = json_str.count('}')
                open_brackets = json_str.count('[')
                close_brackets = json_str.count(']')
                
                # Add missing closures
                json_str = json_str.rstrip()
                if json_str.endswith(','):
                    json_str = json_str[:-1]
                
                # Add missing brackets first, then braces
                json_str += ']' * (open_brackets - close_brackets)
                json_str += '}' * (open_braces - close_braces)
                
                # Clean trailing commas after adding closures
                json_str = re.sub(r',\s*}', '}', json_str)
                json_str = re.sub(r',\s*]', ']', json_str)
                
                return json.loads(json_str)
        except (json.JSONDecodeError, Exception) as e:
            logger.debug(f"Strategy 4 failed: {e}")
        
        # Strategy 5: Extract key-value pairs manually for simple responses
        try:
            # Look for common field patterns
            summary_match = re.search(r'"document_summary"\s*:\s*"([^"]*)"', response)
            if summary_match:
                return {
                    "document_summary": summary_match.group(1),
                    "document_type": "document",
                    "key_subjects": [],
                    "_parsing_note": "Partial extraction from malformed JSON"
                }
        except Exception as e:
            logger.debug(f"Strategy 5 failed: {e}")
        
        # Log a sample of what we couldn't parse
        logger.debug(f"All JSON extraction strategies failed. Response sample: {response[:300]}...")
        
        return None
    
    def _create_fallback_analysis(self, content: str, source_file: str, llm_response: str) -> dict:
        """Create a structured analysis from content when JSON parsing fails."""
        # Extract what we can from the raw content using regex
        import re
        
        # Try to extract any useful information from the LLM's narrative response
        summary = llm_response[:500] if llm_response and llm_response != "Empty LLM response" else "Document analysis - see extracted data"
        
        # Basic extraction from original content
        names = list(set(re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b', content[:5000])))[:10]
        dates = list(set(re.findall(r'\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b', content, re.I)))[:10]
        
        return {
            "document_summary": summary,
            "document_type": self._guess_document_type(source_file, content),
            "people": names[:10],
            "dates": dates[:10],
            "locations": [],
            "key_events": [],
            "key_claims": [],
            "relationships": [],
            "red_flags": [],
            "timeline": [f"{date}: mentioned in document" for date in dates[:5]],
            "investigative_leads": [],
            "questions": [],
            "intelligence_value": "medium",
            "confidence": 0.5,
            "_parsing_note": "JSON parsing failed - used fallback extraction"
        }
    
    def _guess_document_type(self, source_file: str, content: str) -> str:
        """Guess the document type based on filename and content."""
        source_lower = source_file.lower()
        content_lower = content[:1000].lower()
        
        if 'transcript' in source_lower or 'depo' in source_lower:
            return "deposition transcript"
        elif 'photo' in source_lower or 'exhibit' in source_lower:
            return "photo/exhibit"
        elif 'report' in source_lower or 'incident' in source_lower:
            return "incident report"
        elif 'email' in source_lower or 'from:' in content_lower:
            return "email"
        elif 'invoice' in source_lower or 'receipt' in source_lower:
            return "financial document"
        elif 'activity' in source_lower or 'printout' in source_lower:
            return "activity log"
        else:
            return "document"
    
    def synthesize_investigation(self, all_findings: List[dict], all_entities: dict, document_summaries: List[str]) -> dict:
        """Synthesize all collected data into a cohesive investigative narrative.
        
        This method takes all the individual analysis results and creates a
        unified intelligence assessment with connections mapped across sources.
        
        Args:
            all_findings: List of findings from various analysis modules
            all_entities: Aggregated entities from all sources
            document_summaries: Summaries of all analyzed documents
        
        Returns:
            Comprehensive investigation synthesis
        """
        findings_text = "\n".join([
            f"- [{f.get('severity', 'INFO')}] {f.get('title', 'Unknown')}: {f.get('description', '')[:200]}"
            for f in all_findings[:30]
        ])
        
        entities_text = "\n".join([
            f"- {category}: {', '.join(items[:10])}"
            for category, items in all_entities.items()
            if items
        ])
        
        summaries_text = "\n---\n".join(document_summaries[:10])
        
        task_prompt = f"""{self.system_prompt}

SPECIFIC TASK: Investigation Synthesis and Intelligence Assessment

You are synthesizing all collected intelligence into a comprehensive assessment.
Your task is to:

1. CONNECT THE DOTS: Map relationships across all data sources
2. BUILD THE NARRATIVE: What story does the evidence tell?
3. IDENTIFY PATTERNS: What recurring themes emerge across sources?
4. ASSESS THREATS: What risks or concerns are indicated?
5. EVALUATE CONFIDENCE: How reliable is this assessment?
6. IDENTIFY GAPS: What critical information is missing?
7. RECOMMEND ACTIONS: What should be done next?

FINDINGS FROM ANALYSIS:
{findings_text}

ENTITIES IDENTIFIED:
{entities_text}

DOCUMENT SUMMARIES:
{summaries_text}

OUTPUT FORMAT (JSON):
{{
    "executive_summary": "High-level overview for senior review",
    "narrative_assessment": "The story told by the evidence",
    "key_actors": [
        {{"name": "", "role_in_investigation": "", "threat_level": "", "known_connections": []}}
    ],
    "central_events": [
        {{"event": "", "significance": "", "actors_involved": [], "evidence_strength": ""}}
    ],
    "pattern_analysis": [
        {{"pattern": "", "occurrences": 0, "significance": "", "sources": []}}
    ],
    "relationship_network": [
        {{"from": "", "to": "", "relationship": "", "strength": "", "evidence": ""}}
    ],
    "threat_assessment": {{
        "overall_level": "critical/high/medium/low",
        "threat_actors": [],
        "threat_vectors": [],
        "potential_impacts": [],
        "trajectory": "escalating/stable/de-escalating"
    }},
    "timeline_reconstruction": [
        {{"date": "", "event": "", "significance": ""}}
    ],
    "competing_hypotheses": [
        {{"hypothesis": "", "supporting_evidence": [], "contradicting_evidence": [], "probability": 0.0-1.0}}
    ],
    "confidence_assessment": {{
        "overall": 0.0-1.0,
        "factors_supporting": [],
        "factors_limiting": []
    }},
    "intelligence_gaps": [
        {{"gap": "", "impact": "", "collection_recommendation": ""}}
    ],
    "immediate_actions": [],
    "further_investigation_required": [],
    "analyst_notes": ""
}}

CRITICAL: Your response MUST be ONLY the JSON object above. Do not include any text before or after the JSON.
Start directly with {{ and end with }}. If you cannot determine certain information, use empty arrays [] or empty strings ""."""
        
        response = self.generate(task_prompt)
        
        if not response or not response.strip():
            logger.warning("Empty response from LLM for investigation synthesis")
            return {"executive_summary": "Synthesis failed - empty response", "narrative_assessment": ""}
        
        result = self._extract_json_from_response(response)
        if result:
            return result
        
        logger.warning("Failed to parse investigation synthesis")
        return {"executive_summary": response[:500] if response else "Synthesis failed", "narrative_assessment": response if response else ""}
        
        return {"indicator_type": "unknown", "severity": "info"}


# =============================================================================
# LANGCHAIN ENHANCED ANALYZER
# =============================================================================

class LangChainAnalyzer:
    """
    LangChain-enhanced document analysis with advanced chunking and chain-of-thought.
    Uses local Ollama for all LLM operations - no external API calls.
    Optimized for GTX 1070 with safe memory limits.
    """
    
    def __init__(self, model: str = "wizardlm2:latest"):
        self.model = model
        self.available = False
        self.llm = None
        self.text_splitter = None
        
        if LANGCHAIN_AVAILABLE:
            try:
                # Initialize Ollama through LangChain (local only)
                self.llm = Ollama(
                    model=model,
                    base_url="http://localhost:11434",
                    temperature=0.1,
                    num_ctx=8192,    # 8K context - stable for GTX 1070
                    num_predict=6000, # Max output for complete responses
                    stop=[],          # Don't stop early
                )
                
                # Text splitter with smaller chunks for GTX 1070 memory limits
                self.text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=6000,      # Smaller chunks for 8K context
                    chunk_overlap=1000,   # Overlap for context preservation
                    length_function=len,
                    separators=["\n\n", "\n", ". ", " ", ""]
                )
                
                self.available = True
                logger.info(f"LangChain initialized with {model}")
            except Exception as e:
                logger.warning(f"LangChain initialization failed: {e}")
                self.available = False
        else:
            logger.debug("LangChain not available - using standard Ollama")
    
    def analyze_document_with_chain(self, content: str, source_file: str = "") -> dict:
        """
        Analyze document using LangChain with chunking and chain-of-thought reasoning.
        Processes large documents in chunks with context preservation.
        """
        if not self.available or not self.llm:
            return {}
        
        try:
            # Split document into chunks with overlap
            chunks = self.text_splitter.split_text(content)
            
            if len(chunks) == 1:
                # Document fits in single context - direct analysis
                return self._analyze_single_chunk(chunks[0], source_file)
            else:
                # Multi-chunk analysis with synthesis
                logger.info(f"  📑 Document split into {len(chunks)} chunks for analysis")
                return self._analyze_multiple_chunks(chunks, source_file)
                
        except Exception as e:
            logger.error(f"LangChain analysis failed: {e}")
            return {}
    
    def _analyze_single_chunk(self, content: str, source_file: str) -> dict:
        """Analyze a single chunk of content."""
        prompt_template = PromptTemplate(
            input_variables=["content", "source"],
            template=INTELLIGENCE_SYSTEM_PROMPT + """

TASK: Comprehensive Document Analysis

Document Source: {source}

Analyze this document thoroughly. Provide complete analysis without truncation.

DOCUMENT CONTENT:
{content}

REQUIRED OUTPUT FORMAT (JSON):
{{
    "objective": "What this document analysis aims to establish",
    "document_summary": "Comprehensive summary of document contents",
    "document_type": "Type of document",
    "key_findings": [
        "Finding 1 with full explanation",
        "Finding 2 with full explanation"
    ],
    "analysis": {{
        "behavioral_indicators": ["indicator1", "indicator2"],
        "temporal_patterns": ["pattern1", "pattern2"],
        "relationship_mapping": ["relationship1", "relationship2"]
    }},
    "people": ["person1", "person2"],
    "dates": ["date1", "date2"],
    "locations": ["location1", "location2"],
    "key_events": ["event1 with context", "event2 with context"],
    "investigative_insights": [
        "Insight 1 with reasoning",
        "Insight 2 with reasoning"
    ],
    "questions_for_investigation": [
        "Question 1",
        "Question 2"
    ],
    "recommendations": [
        "Recommendation 1",
        "Recommendation 2"
    ],
    "red_flags": ["flag1", "flag2"],
    "intelligence_value": "high/medium/low",
    "confidence": 0.85
}}

Respond with ONLY the JSON object. Be thorough and complete."""
        )
        
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        response = chain.run(content=content, source=source_file)
        
        return self._parse_response(response)
    
    def _analyze_multiple_chunks(self, chunks: List[str], source_file: str) -> dict:
        """Analyze multiple chunks and synthesize findings."""
        chunk_analyses = []
        
        # Analyze each chunk
        for i, chunk in enumerate(chunks):
            logger.info(f"    Analyzing chunk {i+1}/{len(chunks)}...")
            
            chunk_prompt = PromptTemplate(
                input_variables=["chunk_num", "total_chunks", "content", "source"],
                template=INTELLIGENCE_SYSTEM_PROMPT + """

TASK: Document Chunk Analysis (Part {chunk_num} of {total_chunks})

Document Source: {source}

Analyze this section of a larger document. Extract all relevant intelligence.

CHUNK CONTENT:
{content}

Provide analysis as JSON with: key_findings, people, dates, events, insights, questions.
Be thorough - this is part of a larger analysis."""
            )
            
            chain = LLMChain(llm=self.llm, prompt=chunk_prompt)
            response = chain.run(
                chunk_num=i+1,
                total_chunks=len(chunks),
                content=chunk,
                source=source_file
            )
            
            parsed = self._parse_response(response)
            if parsed:
                chunk_analyses.append(parsed)
        
        # Synthesize all chunk analyses
        return self._synthesize_chunks(chunk_analyses, source_file)
    
    def _synthesize_chunks(self, chunk_analyses: List[dict], source_file: str) -> dict:
        """Synthesize multiple chunk analyses into a coherent report."""
        if not chunk_analyses:
            return {}
        
        # Aggregate findings from all chunks
        all_findings = []
        all_people = set()
        all_dates = set()
        all_events = []
        all_insights = []
        all_questions = []
        all_red_flags = []
        
        for analysis in chunk_analyses:
            all_findings.extend(analysis.get('key_findings', []))
            all_people.update(analysis.get('people', []))
            all_dates.update(analysis.get('dates', []))
            all_events.extend(analysis.get('key_events', []))
            all_insights.extend(analysis.get('investigative_insights', analysis.get('insights', [])))
            all_questions.extend(analysis.get('questions_for_investigation', analysis.get('questions', [])))
            all_red_flags.extend(analysis.get('red_flags', []))
        
        # Create synthesis prompt
        synthesis_prompt = PromptTemplate(
            input_variables=["findings", "people", "events", "source"],
            template=INTELLIGENCE_SYSTEM_PROMPT + """

TASK: Intelligence Synthesis

Synthesize analysis from multiple document sections into a cohesive intelligence report.

Document Source: {source}

AGGREGATED FINDINGS:
{findings}

IDENTIFIED PEOPLE: {people}

KEY EVENTS:
{events}

Create a comprehensive synthesis that:
1. Identifies overarching patterns
2. Connects related findings
3. Highlights critical intelligence
4. Provides actionable recommendations

OUTPUT FORMAT (JSON):
{{
    "objective": "Synthesis objective",
    "document_summary": "Comprehensive synthesis of all sections",
    "document_type": "document type",
    "key_findings": ["synthesized finding 1", "synthesized finding 2"],
    "analysis": {{
        "overarching_patterns": ["pattern1", "pattern2"],
        "cross_references": ["connection1", "connection2"],
        "critical_intelligence": ["intel1", "intel2"]
    }},
    "people": ["all identified people"],
    "dates": ["all dates"],
    "locations": ["all locations"],
    "key_events": ["all events with context"],
    "investigative_insights": ["synthesized insights"],
    "questions_for_investigation": ["prioritized questions"],
    "recommendations": ["actionable recommendations"],
    "red_flags": ["consolidated red flags"],
    "intelligence_value": "high/medium/low",
    "confidence": 0.85
}}

Respond with ONLY the JSON object."""
        )
        
        chain = LLMChain(llm=self.llm, prompt=synthesis_prompt)
        response = chain.run(
            findings="\n".join(f"- {f}" for f in all_findings[:20]),
            people=", ".join(list(all_people)[:15]),
            events="\n".join(f"- {e}" for e in all_events[:15]),
            source=source_file
        )
        
        result = self._parse_response(response)
        
        # Ensure all extracted data is preserved
        if result:
            result['people'] = list(all_people)
            result['dates'] = list(all_dates)
            result['questions_for_investigation'] = list(set(all_questions))
            result['red_flags'] = list(set(all_red_flags))
        
        return result
    
    def _parse_response(self, response: str) -> dict:
        """Parse LLM response to extract JSON."""
        if not response:
            return {}
        
        try:
            # Find JSON in response
            start = response.find('{')
            end = response.rfind('}') + 1
            if start != -1 and end > start:
                import json
                return json.loads(response[start:end])
        except json.JSONDecodeError:
            pass
        
        # Try to fix common JSON issues
        try:
            import re
            json_str = response[response.find('{'):response.rfind('}')+1]
            json_str = re.sub(r',\s*}', '}', json_str)
            json_str = re.sub(r',\s*]', ']', json_str)
            return json.loads(json_str)
        except:
            pass
        
        return {}


# =============================================================================
# ENHANCED ANALYSIS MODULE WITH LLM
# =============================================================================

class LLMEnhancedAnalysisModule(AnalysisModule):
    """Analysis module enhanced with local LLM capabilities.
    
    This module performs comprehensive document analysis including:
    - Semantic understanding of document content
    - Entity extraction with context
    - Relationship mapping across documents
    - Pattern detection and anomaly identification
    - Progressive analysis that builds on previous findings
    - Investigation synthesis across all sources
    
    Uses LangChain when available for advanced chunking and chain-of-thought.
    Optimized for GTX 1070 with WizardLM2.
    """
    
    def __init__(self, model: str = "wizardlm2:latest"):
        self.llm = OllamaAnalyzer(model=model)
        self.langchain_analyzer = LangChainAnalyzer(model=model) if LANGCHAIN_AVAILABLE else None
        self.document_analyses = []  # Store individual document analyses
        self.all_entities = {}  # Aggregated entities
        self.previous_context = ""  # Context from previous runs
        
        if self.langchain_analyzer and self.langchain_analyzer.available:
            logger.info("📦 LangChain enhanced analysis enabled")
        else:
            logger.info("📦 Using standard Ollama analysis")
    
    def analyze(
        self, 
        data: List[Dict[str, Any]], 
        config: InvestigationConfig
    ) -> List[Finding]:
        """Analyze data using LLM for comprehensive investigative insights.
        
        This method:
        1. Analyzes each document for content and meaning
        2. Extracts entities and relationships
        3. Identifies patterns and anomalies
        4. Builds connections across all documents
        5. Synthesizes findings into cohesive intelligence
        """
        findings = []
        
        # Re-check LLM connection before analysis
        if not self.llm.available:
            self.llm.check_connection()
        
        if not self.llm.available:
            logger.warning("LLM not available, skipping LLM-enhanced analysis")
            return findings
        
        # Load previous context if available
        self._load_previous_context(config)
        
        # Group records by source file for document-level analysis
        documents_by_source = self._group_by_source(data)
        
        logger.info(f"📄 Analyzing {len(documents_by_source)} document(s) with LLM...")
        
        # Phase 1: Individual Document Analysis
        document_summaries = []
        for source_file, records in documents_by_source.items():
            doc_findings = self._analyze_document(source_file, records, config)
            findings.extend(doc_findings)
            
            # Collect document summary for synthesis
            if self.document_analyses:
                last_analysis = self.document_analyses[-1]
                if last_analysis.get('document_summary'):
                    document_summaries.append(
                        f"[{source_file}]: {last_analysis.get('document_summary', '')}"
                    )
        
        # Phase 2: Cross-Document Entity Analysis
        all_text = self._extract_all_text(data)
        if all_text:
            entity_findings = self._analyze_entities_comprehensive(all_text)
            findings.extend(entity_findings)
        
        # Phase 3: Log and Anomaly Detection
        log_findings = self._analyze_logs(data)
        findings.extend(log_findings)
        
        # Phase 4: Timeline and Behavioral Analysis
        temporal_findings = self._analyze_temporal_patterns(data)
        findings.extend(temporal_findings)
        
        # Phase 5: High-Value Data Deep Dive
        sensitive_findings = self._analyze_sensitive_data(data)
        findings.extend(sensitive_findings)
        
        # Phase 6: Investigation Synthesis (if we have enough data)
        if len(findings) >= 3 or len(document_summaries) >= 2:
            synthesis_findings = self._synthesize_investigation(findings, document_summaries)
            findings.extend(synthesis_findings)
        
        logger.info(f"✅ LLM analysis complete: {len(findings)} findings generated")
        
        return findings
    
    def _load_previous_context(self, config: InvestigationConfig) -> None:
        """Load context from previous analysis runs."""
        report_json = config.output_dir / 'report.json'
        if report_json.exists():
            try:
                with report_json.open('r', encoding='utf-8') as f:
                    prev_data = json.load(f)
                    
                # Build context summary
                context_parts = []
                
                # Previous findings summary
                prev_findings = prev_data.get('findings', [])
                if prev_findings:
                    context_parts.append(f"Previous analysis found {len(prev_findings)} findings:")
                    for f in prev_findings[:5]:
                        context_parts.append(f"  - [{f.get('severity')}] {f.get('title')}")
                
                # Previous entities
                prev_entities = prev_data.get('entities_summary', {})
                if prev_entities:
                    context_parts.append("\nPreviously identified entities:")
                    for etype, values in prev_entities.items():
                        if values:
                            context_parts.append(f"  - {etype}: {', '.join(values[:5])}")
                
                self.previous_context = "\n".join(context_parts)
                logger.info(f"📚 Loaded context from previous run ({len(prev_findings)} findings)")
                
            except Exception as e:
                logger.debug(f"Could not load previous context: {e}")
    
    def _group_by_source(self, data: List[Dict[str, Any]]) -> Dict[str, List[Dict]]:
        """Group records by their source file, preserving subdirectory structure."""
        grouped = {}
        for record in data:
            source = record.get('source', 'unknown')
            if source != 'unknown':
                # Try to get a meaningful path (with subdirectory if present)
                source_path = Path(source)
                # Get the last 2-3 parts of the path to show subdirectory context
                parts = source_path.parts
                if len(parts) >= 3:
                    # Include parent folder and filename: "subfolder/file.txt"
                    source_name = str(Path(*parts[-2:]))
                elif len(parts) >= 2:
                    source_name = str(Path(*parts[-2:]))
                else:
                    source_name = source_path.name
            else:
                source_name = 'unknown'
            
            if source_name not in grouped:
                grouped[source_name] = []
            grouped[source_name].append(record)
        return grouped
    
    def _extract_all_text(self, data: List[Dict[str, Any]]) -> str:
        """Extract all text content from records."""
        text_parts = []
        for record in data:
            record_data = record.get('data', {})
            if isinstance(record_data, dict):
                # Extract text fields
                for key in ['content', 'text', 'body', 'message', 'description', 'summary']:
                    if key in record_data and record_data[key]:
                        text_parts.append(str(record_data[key]))
                # Also include the full record if it's small
                record_str = safe_json_dumps(record_data)
                if len(record_str) < 2000:
                    text_parts.append(record_str)
            elif isinstance(record_data, str):
                text_parts.append(record_data)
        
        return "\n\n---\n\n".join(text_parts[:30])  # Limit for context window
    
    def _analyze_document(self, source_file: str, records: List[Dict], config: InvestigationConfig) -> List[Finding]:
        """Perform comprehensive analysis of a single document.
        
        Uses LangChain when available for enhanced chunking and chain-of-thought.
        Falls back to standard Ollama analysis otherwise.
        """
        findings = []
        
        # Combine all content from this document
        content_parts = []
        for record in records:
            record_data = record.get('data', {})
            if isinstance(record_data, dict):
                content_parts.append(safe_json_dumps(record_data))
            elif isinstance(record_data, str):
                content_parts.append(record_data)
        
        if not content_parts:
            return findings
        
        combined_content = "\n".join(content_parts)
        
        # Skip very small documents
        if len(combined_content) < 50:
            return findings
        
        logger.info(f"  📖 Analyzing: {source_file}")
        
        # Use LangChain if available for better chunking and analysis
        if self.langchain_analyzer and self.langchain_analyzer.available:
            analysis = self.langchain_analyzer.analyze_document_with_chain(
                content=combined_content,
                source_file=source_file
            )
            if not analysis:
                # Fallback to standard analysis
                analysis = self.llm.analyze_document_content(
                    content=combined_content,
                    source_file=source_file,
                    previous_context=self.previous_context
                )
        else:
            # Standard Ollama analysis
            analysis = self.llm.analyze_document_content(
                content=combined_content,
                source_file=source_file,
                previous_context=self.previous_context
            )
        
        if not analysis or analysis.get('document_summary') == 'Analysis failed':
            return findings
        
        # Store for synthesis
        self.document_analyses.append(analysis)
        
        # Create findings from document analysis
        
        # 1. Document Summary Finding
        findings.append(Finding(
            id=str(uuid.uuid4()),
            severity=self._assess_severity(analysis.get('intelligence_value', 'medium')),
            title=f"Document Analysis: {source_file}",
            description=analysis.get('document_summary', 'Document analyzed'),
            entities=[],
            evidence=[
                f"Document Type: {analysis.get('document_type', 'Unknown')}",
                f"Intelligence Value: {analysis.get('intelligence_value', 'Unknown')}",
                f"Confidence: {analysis.get('confidence', analysis.get('confidence_assessment', 'N/A'))}",
            ],
            timestamp=datetime.datetime.now(),
            recommendations=analysis.get('recommended_next_steps', [])[:3]
        ))
        
        # 2. Key Subjects Finding - handle both old and new schema
        subjects = analysis.get('key_subjects', [])
        people = analysis.get('people', [])
        if subjects:
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.MEDIUM,
                title=f"Key Subjects Identified: {len(subjects)} person(s)/entity(ies)",
                description="Significant actors identified in document analysis",
                entities=[],
                evidence=[
                    f"{s.get('name', 'Unknown')}: {s.get('role', 'Unknown role')} - {s.get('significance', '')}"
                    for s in subjects[:10]
                ] if isinstance(subjects[0], dict) else subjects[:10],
                timestamp=datetime.datetime.now(),
                recommendations=["Verify identities", "Map connections to other entities"]
            ))
        elif people:
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.MEDIUM,
                title=f"People Identified: {len(people)} person(s)",
                description="People mentioned in document",
                entities=[],
                evidence=people[:10] if isinstance(people, list) else [str(people)],
                timestamp=datetime.datetime.now(),
                recommendations=["Verify identities", "Map connections to other entities"]
            ))
        
        # 3. Key Events Finding - handle both schemas
        events = analysis.get('key_events', [])
        if events:
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.MEDIUM,
                title=f"Key Events Identified: {len(events)} event(s)",
                description="Significant events extracted from document",
                entities=[],
                evidence=[
                    f"{e.get('date', 'Unknown date')}: {e.get('event', 'Unknown')} - {e.get('significance', '')}"
                    for e in events[:10]
                ] if events and isinstance(events[0], dict) else events[:10],
                timestamp=datetime.datetime.now(),
                recommendations=["Verify event details", "Establish timeline"]
            ))
        
        # 4. Relationships Finding - handle both schemas
        relationships = analysis.get('relationships_identified', []) or analysis.get('relationships', [])
        if relationships:
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.MEDIUM,
                title=f"Relationships Mapped: {len(relationships)} connection(s)",
                description="Relationships between entities identified",
                entities=[],
                evidence=[
                    f"{r.get('entity1', '?')} <-> {r.get('entity2', '?')}: {r.get('relationship_type', 'Unknown')}"
                    for r in relationships[:10]
                ] if relationships and isinstance(relationships[0], dict) else relationships[:10],
                timestamp=datetime.datetime.now(),
                recommendations=["Visualize relationship network", "Identify central nodes"]
            ))
        
        # 5. Red Flags Finding - handle both schemas
        red_flags = analysis.get('red_flags', [])
        if red_flags:
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.HIGH,
                title=f"⚠️ Red Flags Detected: {len(red_flags)} indicator(s)",
                description="Warning indicators requiring immediate attention",
                entities=[],
                evidence=[
                    f"{rf.get('indicator', rf) if isinstance(rf, dict) else rf}: {rf.get('significance', '') if isinstance(rf, dict) else ''}"
                    for rf in red_flags[:10]
                ] if red_flags else [],
                timestamp=datetime.datetime.now(),
                recommendations=[rf.get('recommended_action', '') for rf in red_flags[:5] if isinstance(rf, dict) and rf.get('recommended_action')] or ["Review flagged items"]
            ))
        
        # 6. Investigative Leads Finding - handle both schemas
        leads = analysis.get('investigative_leads', [])
        if leads:
            if leads and isinstance(leads[0], dict):
                high_priority_leads = [l for l in leads if l.get('priority') == 'high']
                evidence = [
                    f"[{l.get('priority', 'medium').upper()}] {l.get('lead', 'Unknown')}"
                    for l in leads[:10]
                ]
                recs = [l.get('lead', '') for l in high_priority_leads[:5]]
            else:
                high_priority_leads = []
                evidence = leads[:10]
                recs = leads[:5]
            
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.INFO if not high_priority_leads else SeverityLevel.MEDIUM,
                title=f"Investigative Leads: {len(leads)} lead(s) generated",
                description="Actionable investigative directions identified",
                entities=[],
                evidence=evidence,
                timestamp=datetime.datetime.now(),
                recommendations=recs
            ))
        
        # 7. Questions Raised - handle both schemas
        questions = analysis.get('questions_raised', []) or analysis.get('questions', [])
        if questions:
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.INFO,
                title=f"Questions for Investigation: {len(questions)} question(s)",
                description="Unanswered questions requiring further investigation",
                entities=[],
                evidence=questions[:10],
                timestamp=datetime.datetime.now(),
                recommendations=["Address each question systematically"]
            ))
        
        # Aggregate entities for cross-document analysis
        entities_found = analysis.get('entities_identified', {})
        if isinstance(entities_found, dict):
            for category, items in entities_found.items():
                if items:
                    if category not in self.all_entities:
                        self.all_entities[category] = []
                    self.all_entities[category].extend(items if isinstance(items, list) else [items])
        
        return findings
    
    def _assess_severity(self, intelligence_value: str) -> SeverityLevel:
        """Convert intelligence value to severity level."""
        value_map = {
            'critical': SeverityLevel.CRITICAL,
            'high': SeverityLevel.HIGH,
            'medium': SeverityLevel.MEDIUM,
            'low': SeverityLevel.LOW,
        }
        return value_map.get(intelligence_value.lower(), SeverityLevel.INFO)
    
    def _analyze_entities_comprehensive(self, text: str) -> List[Finding]:
        """Perform comprehensive entity extraction across all text."""
        findings = []
        
        llm_entities = self.llm.analyze_entities(text)
        
        if llm_entities:
            entity_count = sum(len(v) for v in llm_entities.values() if isinstance(v, list))
            if entity_count > 0:
                # Merge with aggregated entities
                for category, items in llm_entities.items():
                    if isinstance(items, list) and items:
                        if category not in self.all_entities:
                            self.all_entities[category] = []
                        self.all_entities[category].extend(items)
                
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.INFO,
                    title=f"Entity Extraction: {entity_count} entities identified",
                    description="Comprehensive entity extraction across all documents",
                    entities=[],
                    evidence=[
                        f"{k}: {', '.join(v[:5])}{' (+more)' if len(v) > 5 else ''}" 
                        for k, v in llm_entities.items() 
                        if isinstance(v, list) and v
                    ],
                    timestamp=datetime.datetime.now(),
                    recommendations=["Cross-reference entities across sources", "Verify entity identities"]
                ))
        
        return findings
    
    def _analyze_logs(self, data: List[Dict[str, Any]]) -> List[Finding]:
        """Analyze log entries for anomalies."""
        findings = []
        
        log_data = []
        for record in data:
            record_data = record.get('data', {})
            if record.get('type') == 'parsed_record' and isinstance(record_data, dict):
                if 'message' in record_data:
                    log_data.append(str(record_data.get('message', '')))
        
        if not log_data:
            return findings
        
        anomaly_result = self.llm.analyze_log_anomaly(log_data)
        
        if anomaly_result.get('risk_level') not in ['unknown', 'low', None]:
            severity_map = {
                'critical': SeverityLevel.CRITICAL,
                'high': SeverityLevel.HIGH,
                'medium': SeverityLevel.MEDIUM,
            }
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=severity_map.get(
                    anomaly_result.get('risk_level', 'medium'),
                    SeverityLevel.MEDIUM
                ),
                title=f"Log Anomaly Detection: {anomaly_result.get('risk_level', 'unknown')} risk",
                description="Anomalies detected through behavioral log analysis",
                entities=[],
                evidence=[
                    f.get('description', '') 
                    for f in anomaly_result.get('findings', [])[:5]
                ],
                timestamp=datetime.datetime.now(),
                recommendations=anomaly_result.get('recommendations', [])[:5]
            ))
            
            iocs = anomaly_result.get('iocs', [])
            if iocs:
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.HIGH,
                    title=f"IOCs Identified: {len(iocs)} indicator(s)",
                    description="Indicators of Compromise detected",
                    entities=[],
                    evidence=iocs[:10],
                    timestamp=datetime.datetime.now(),
                    recommendations=["Block/investigate identified IOCs"]
                ))
        
        return findings
    
    def _analyze_temporal_patterns(self, data: List[Dict[str, Any]]) -> List[Finding]:
        """Analyze temporal patterns and timeline."""
        findings = []
        
        timeline_events = []
        for record in data:
            record_data = record.get('data', {})
            if isinstance(record_data, dict):
                timestamp = record_data.get('timestamp') or record_data.get('time') or record_data.get('date')
                if timestamp:
                    timeline_events.append({
                        'timestamp': str(timestamp),
                        'description': record_data.get('message') or record_data.get('event') or str(record_data)[:200]
                    })
        
        if len(timeline_events) < 3:
            return findings
        
        timeline_result = self.llm.analyze_timeline(timeline_events)
        
        if timeline_result and timeline_result.get('timeline_summary'):
            key_events = timeline_result.get('key_events', [])
            anomalies = timeline_result.get('anomalies', [])
            
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.MEDIUM if anomalies else SeverityLevel.INFO,
                title=f"Timeline Analysis: {len(key_events)} key events",
                description=timeline_result.get('timeline_summary', 'Timeline reconstructed'),
                entities=[],
                evidence=[
                    f"Timespan: {timeline_result.get('total_timespan', 'Unknown')}",
                    *[f"Event: {e.get('event', '')} at {e.get('timestamp', '')}" for e in key_events[:5]],
                    *[f"Anomaly: {a.get('description', '')}" for a in anomalies[:3]]
                ],
                timestamp=datetime.datetime.now(),
                recommendations=timeline_result.get('investigative_recommendations', [])[:5]
            ))
            
            phases = timeline_result.get('behavioral_phases', [])
            if phases:
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.HIGH,
                    title=f"Behavioral Phases: {len(phases)} phase(s) detected",
                    description="Activity phases suggesting coordinated behavior",
                    entities=[],
                    evidence=[f"Phase '{p.get('phase')}': {p.get('start')} - {p.get('end')}" for p in phases[:5]],
                    timestamp=datetime.datetime.now(),
                    recommendations=["Review phases for operational patterns"]
                ))
        
        return findings
    
    def _analyze_sensitive_data(self, data: List[Dict[str, Any]]) -> List[Finding]:
        """Analyze high-value/sensitive data."""
        findings = []
        
        sensitive_indicators = [
            'password', 'credential', 'ssn', 'social security', 'credit', 'bank',
            'secret', 'key', 'token', 'auth', 'account', 'routing', 'wire',
            'payment', 'bitcoin', 'wallet', 'private', 'confidential'
        ]
        
        high_value_data = []
        for record in data:
            record_data = record.get('data', {})
            record_str = str(record_data).lower()
            if any(indicator in record_str for indicator in sensitive_indicators):
                high_value_data.append(record_data)
        
        if not high_value_data:
            return findings
        
        combined = "\n".join([
            safe_json_dumps(d) if isinstance(d, dict) else str(d) 
            for d in high_value_data[:10]
        ])
        
        result = self.llm.deep_dive_analysis(combined, context="Sensitive/high-value data analysis")
        
        if result and result.get('key_findings'):
            key_findings = result.get('key_findings', [])
            risk = result.get('risk_assessment', {})
            
            risk_level = risk.get('level', 'unknown').lower()
            severity_map = {
                'critical': SeverityLevel.CRITICAL, 
                'high': SeverityLevel.HIGH,
                'medium': SeverityLevel.MEDIUM, 
                'low': SeverityLevel.LOW
            }
            
            findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=severity_map.get(risk_level, SeverityLevel.MEDIUM),
                title=f"Sensitive Data Analysis: {len(key_findings)} finding(s)",
                description=result.get('objective', 'Analysis of sensitive data'),
                entities=[],
                evidence=[
                    *[f"{f.get('finding', '')} (confidence: {f.get('confidence', 'N/A')})" for f in key_findings[:5]],
                    f"Risk Level: {risk_level}",
                    *risk.get('factors', [])[:3]
                ],
                timestamp=datetime.datetime.now(),
                recommendations=result.get('recommended_next_steps', [])[:5]
            ))
            
            gaps = result.get('intelligence_gaps', [])
            if gaps:
                findings.append(Finding(
                    id=str(uuid.uuid4()),
                    severity=SeverityLevel.INFO,
                    title=f"Intelligence Gaps: {len(gaps)} area(s)",
                    description="Areas requiring additional investigation",
                    entities=[],
                    evidence=gaps[:10],
                    timestamp=datetime.datetime.now(),
                    recommendations=result.get('questions_for_investigation', [])[:5]
                ))
        
        return findings
    
    def _synthesize_investigation(self, findings: List[Finding], document_summaries: List[str]) -> List[Finding]:
        """Synthesize all findings into cohesive intelligence assessment."""
        synthesis_findings = []
        
        # Prepare findings for synthesis
        findings_data = [
            {
                'severity': f.severity.name,
                'title': f.title,
                'description': f.description[:200]
            }
            for f in findings[:30]
        ]
        
        synthesis = self.llm.synthesize_investigation(
            all_findings=findings_data,
            all_entities=self.all_entities,
            document_summaries=document_summaries
        )
        
        if not synthesis or synthesis.get('executive_summary') == 'Synthesis failed':
            return synthesis_findings
        
        # Create synthesis finding
        threat = synthesis.get('threat_assessment', {})
        threat_level = threat.get('overall_level', 'unknown').lower()
        
        severity_map = {
            'critical': SeverityLevel.CRITICAL,
            'high': SeverityLevel.HIGH,
            'medium': SeverityLevel.MEDIUM,
            'low': SeverityLevel.LOW
        }
        
        synthesis_findings.append(Finding(
            id=str(uuid.uuid4()),
            severity=severity_map.get(threat_level, SeverityLevel.MEDIUM),
            title="🔍 INVESTIGATION SYNTHESIS",
            description=synthesis.get('executive_summary', 'Investigation synthesis complete'),
            entities=[],
            evidence=[
                f"Narrative: {synthesis.get('narrative_assessment', 'N/A')[:300]}",
                f"Threat Level: {threat_level.upper()}",
                f"Trajectory: {threat.get('trajectory', 'Unknown')}",
            ],
            timestamp=datetime.datetime.now(),
            recommendations=synthesis.get('immediate_actions', [])[:5]
        ))
        
        # Key actors finding
        actors = synthesis.get('key_actors', [])
        if actors:
            synthesis_findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.HIGH if any(a.get('threat_level') == 'high' for a in actors) else SeverityLevel.MEDIUM,
                title=f"Key Actors Identified: {len(actors)} actor(s)",
                description="Central figures in the investigation",
                entities=[],
                evidence=[
                    f"{a.get('name', 'Unknown')}: {a.get('role_in_investigation', '')} [Threat: {a.get('threat_level', 'Unknown')}]"
                    for a in actors[:10]
                ],
                timestamp=datetime.datetime.now(),
                recommendations=["Deep-dive on high-threat actors", "Map actor networks"]
            ))
        
        # Competing hypotheses
        hypotheses = synthesis.get('competing_hypotheses', [])
        if hypotheses:
            synthesis_findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.INFO,
                title=f"Competing Hypotheses: {len(hypotheses)} hypothesis(es)",
                description="Alternative explanations for observed evidence",
                entities=[],
                evidence=[
                    f"{h.get('hypothesis', 'Unknown')} (probability: {h.get('probability', 'N/A')})"
                    for h in hypotheses[:5]
                ],
                timestamp=datetime.datetime.now(),
                recommendations=["Gather evidence to differentiate hypotheses"]
            ))
        
        # Further investigation required
        further = synthesis.get('further_investigation_required', [])
        if further:
            synthesis_findings.append(Finding(
                id=str(uuid.uuid4()),
                severity=SeverityLevel.MEDIUM,
                title=f"Further Investigation Required: {len(further)} area(s)",
                description="Critical areas requiring additional investigative effort",
                entities=[],
                evidence=further[:10],
                timestamp=datetime.datetime.now(),
                recommendations=further[:5]
            ))
        
        return synthesis_findings


# =============================================================================
# REPORTING ENGINE
# =============================================================================

class ReportGenerator(ABC):
    """Abstract base class for report generators."""
    
    @abstractmethod
    def generate(
        self, 
        investigation: 'Investigation',
        output_path: Path
    ) -> None:
        """Generate investigation report."""
        pass


class MarkdownReportGenerator(ReportGenerator):
    """Generate Markdown reports in intelligence report format."""
    
    def generate(
        self, 
        investigation: 'Investigation',
        output_path: Path
    ) -> None:
        """Generate Markdown report following intelligence report format."""
        with output_path.open('w', encoding='utf-8') as f:
            # Header with classification markings (placeholder)
            f.write("# 🔍 OSINT INTELLIGENCE REPORT\n\n")
            f.write("---\n\n")
            
            # Report metadata
            f.write("## Report Metadata\n\n")
            f.write(f"| Field | Value |\n")
            f.write(f"|-------|-------|\n")
            f.write(f"| **Investigation Name** | {investigation.config.name} |\n")
            f.write(f"| **Investigation Type** | {investigation.config.investigation_type.name} |\n")
            f.write(f"| **Report Generated** | {format_timestamp()} |\n")
            f.write(f"| **Targets** | {', '.join(investigation.config.targets)} |\n")
            f.write(f"| **Data Records Analyzed** | {len(investigation.collected_data)} |\n")
            f.write(f"| **Findings Generated** | {len(investigation.findings)} |\n")
            f.write("\n---\n\n")
            
            # Section 1: Objective
            f.write("## 1. Objective\n\n")
            f.write(f"This investigation was initiated to analyze {investigation.config.investigation_type.name.lower().replace('_', ' ')} ")
            f.write(f"related to the following targets: **{', '.join(investigation.config.targets)}**.\n\n")
            f.write("The analysis leverages multi-discipline intelligence methodologies combining:\n")
            f.write("- **FBI-style** behavioral profiling and threat assessment\n")
            f.write("- **CIA HUMINT** pattern recognition and source evaluation\n")
            f.write("- **NSA OSINT** technical artifact analysis\n")
            f.write("- **AI-enhanced** correlation and anomaly detection\n\n")
            
            # Section 2: Key Findings
            f.write("## 2. Key Findings\n\n")
            severity_counts = Counter(f.severity for f in investigation.findings)
            
            # Summary table
            f.write("### Findings Summary\n\n")
            f.write("| Severity | Count | Status |\n")
            f.write("|----------|-------|--------|\n")
            severity_icons = {
                SeverityLevel.CRITICAL: "🔴",
                SeverityLevel.HIGH: "🟠", 
                SeverityLevel.MEDIUM: "🟡",
                SeverityLevel.LOW: "🟢",
                SeverityLevel.INFO: "🔵"
            }
            for severity in SeverityLevel:
                count = severity_counts.get(severity, 0)
                icon = severity_icons.get(severity, "⚪")
                status = "REQUIRES IMMEDIATE ATTENTION" if severity == SeverityLevel.CRITICAL and count > 0 else "Review Required" if count > 0 else "None"
                f.write(f"| {icon} **{severity.name}** | {count} | {status} |\n")
            f.write("\n")
            
            # Critical and High findings highlighted
            critical_high = [f for f in investigation.findings if f.severity in [SeverityLevel.CRITICAL, SeverityLevel.HIGH]]
            if critical_high:
                f.write("### ⚠️ Priority Findings\n\n")
                for i, finding in enumerate(critical_high[:5], 1):
                    f.write(f"**{i}. {finding.title}** ({finding.severity.name})\n\n")
                    f.write(f"> {finding.description}\n\n")
                f.write("\n")
            
            # Section 3: Analysis
            f.write("## 3. Analysis\n\n")
            f.write("### Detailed Findings\n\n")
            
            for finding in sorted(
                investigation.findings, 
                key=lambda x: x.severity.value, 
                reverse=True
            ):
                severity_icon = severity_icons.get(finding.severity, "⚪")
                f.write(f"#### {severity_icon} {finding.title}\n\n")
                f.write(f"**Severity:** {finding.severity.name} | **ID:** `{finding.id[:8]}`\n\n")
                f.write(f"**Analysis:** {finding.description}\n\n")
                
                if finding.evidence:
                    f.write("**Supporting Evidence:**\n\n")
                    for evidence in finding.evidence:
                        f.write(f"- {evidence}\n")
                    f.write("\n")
                
                if finding.recommendations:
                    f.write("**Recommended Actions:**\n\n")
                    for rec in finding.recommendations:
                        f.write(f"- [ ] {rec}\n")
                    f.write("\n")
                
                f.write("---\n\n")
            
            # Section 4: Investigative Insights
            f.write("## 4. Investigative Insights\n\n")
            
            # Pattern analysis
            finding_types = Counter()
            for finding in investigation.findings:
                if 'entity' in finding.title.lower() or 'extract' in finding.title.lower():
                    finding_types['Entity Extraction'] += 1
                elif 'anomaly' in finding.title.lower() or 'anomalies' in finding.title.lower():
                    finding_types['Anomaly Detection'] += 1
                elif 'timeline' in finding.title.lower() or 'temporal' in finding.title.lower():
                    finding_types['Timeline Analysis'] += 1
                elif 'ioc' in finding.title.lower() or 'indicator' in finding.title.lower():
                    finding_types['IOC Identification'] += 1
                elif 'behavioral' in finding.title.lower() or 'pattern' in finding.title.lower():
                    finding_types['Behavioral Analysis'] += 1
                elif 'deep' in finding.title.lower() or 'intelligence' in finding.title.lower():
                    finding_types['Deep Intelligence'] += 1
                else:
                    finding_types['General Analysis'] += 1
            
            if finding_types:
                f.write("### Analysis Methods Applied\n\n")
                for method, count in finding_types.most_common():
                    f.write(f"- **{method}**: {count} findings generated\n")
                f.write("\n")
            
            # Section 5: Questions for Further Investigation
            f.write("## 5. Questions for Further Investigation\n\n")
            f.write("Based on the analysis, the following questions warrant further investigation:\n\n")
            
            question_count = 1
            for finding in investigation.findings:
                if 'intelligence gap' in finding.title.lower() or finding.evidence:
                    for ev in finding.evidence[:2]:
                        if '?' in str(ev) or 'unknown' in str(ev).lower():
                            f.write(f"{question_count}. {ev}\n")
                            question_count += 1
                            if question_count > 5:
                                break
                if question_count > 5:
                    break
            
            if question_count == 1:
                f.write("1. What is the full scope of affected systems or individuals?\n")
                f.write("2. Are there additional indicators of compromise not yet identified?\n")
                f.write("3. What is the timeline of initial compromise to discovery?\n")
                f.write("4. Are there connections to known threat actors or campaigns?\n")
                f.write("5. What data may have been accessed or exfiltrated?\n")
            f.write("\n")
            
            # Section 6: Recommendations and Next Steps
            f.write("## 6. Recommendations and Next Steps\n\n")
            
            # Collect all unique recommendations
            all_recs = []
            for finding in investigation.findings:
                all_recs.extend(finding.recommendations)
            unique_recs = list(dict.fromkeys(all_recs))  # Preserve order, remove duplicates
            
            if unique_recs:
                f.write("### Immediate Actions\n\n")
                for i, rec in enumerate(unique_recs[:10], 1):
                    f.write(f"{i}. {rec}\n")
                f.write("\n")
            
            f.write("### Strategic Recommendations\n\n")
            f.write("- Conduct follow-up investigation on high-severity findings\n")
            f.write("- Implement monitoring for identified indicators\n")
            f.write("- Document lessons learned and update detection capabilities\n")
            f.write("- Consider engagement with specialized teams if threat escalates\n")
            f.write("\n")
            
            # Appendix: Data Sources
            f.write("---\n\n")
            f.write("## Appendix A: Data Sources\n\n")
            for source in investigation.config.data_sources:
                f.write(f"- `{source}`\n")
            
            # Footer
            f.write("\n---\n\n")
            f.write("*This intelligence report was generated by the OSINT Investigation Framework*\n")
            f.write("*Analysis powered by multi-discipline LLM intelligence synthesis*\n")


class JSONReportGenerator(ReportGenerator):
    """Generate JSON reports."""
    
    def generate(
        self, 
        investigation: 'Investigation',
        output_path: Path
    ) -> None:
        """Generate JSON report."""
        report = {
            'investigation': {
                'name': investigation.config.name,
                'type': investigation.config.investigation_type.name,
                'targets': investigation.config.targets,
                'timestamp': format_timestamp()
            },
            'summary': {
                'total_records': len(investigation.collected_data),
                'total_findings': len(investigation.findings),
                'severity_distribution': {
                    severity.name: sum(
                        1 for f in investigation.findings if f.severity == severity
                    )
                    for severity in SeverityLevel
                }
            },
            'findings': [
                {
                    'id': f.id,
                    'severity': f.severity.name,
                    'title': f.title,
                    'description': f.description,
                    'evidence': f.evidence,
                    'recommendations': f.recommendations,
                    'timestamp': f.timestamp.isoformat()
                }
                for f in investigation.findings
            ],
            'data_sources': [str(s) for s in investigation.config.data_sources]
        }
        
        with output_path.open('w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)


class CSVReportGenerator(ReportGenerator):
    """Generate CSV reports."""
    
    def generate(
        self, 
        investigation: 'Investigation',
        output_path: Path
    ) -> None:
        """Generate CSV report."""
        with output_path.open('w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=[
                'id', 'severity', 'title', 'description', 
                'evidence', 'recommendations', 'timestamp'
            ])
            writer.writeheader()
            
            for finding in investigation.findings:
                writer.writerow({
                    'id': finding.id,
                    'severity': finding.severity.name,
                    'title': finding.title,
                    'description': finding.description,
                    'evidence': '; '.join(finding.evidence),
                    'recommendations': '; '.join(finding.recommendations),
                    'timestamp': finding.timestamp.isoformat()
                })


# =============================================================================
# CACHING SYSTEM
# =============================================================================

class Cache:
    """Simple file-based cache for investigation data."""
    
    def __init__(self, cache_dir: Path, ttl: int = 3600):
        self.cache_dir = cache_dir
        self.ttl = ttl
        self.cache_dir.mkdir(parents=True, exist_ok=True)
    
    def get(self, key: str) -> Optional[Dict[str, Any]]:
        """Get cached value."""
        cache_file = self.cache_dir / f"{self._hash_key(key)}.json"
        
        if cache_file.exists():
            mtime = cache_file.stat().st_mtime
            if datetime.datetime.now().timestamp() - mtime < self.ttl:
                with cache_file.open() as f:
                    return json.load(f)
            else:
                cache_file.unlink()  # Expired
        
        return None
    
    def set(self, key: str, value: Dict[str, Any]) -> None:
        """Set cached value."""
        cache_file = self.cache_dir / f"{self._hash_key(key)}.json"
        with cache_file.open('w') as f:
            json.dump(value, f)
    
    def clear(self) -> None:
        """Clear all cached data."""
        for cache_file in self.cache_dir.glob('*.json'):
            cache_file.unlink()
    
    @staticmethod
    def _hash_key(key: str) -> str:
        """Hash cache key."""
        return hashlib.md5(key.encode()).hexdigest()


# =============================================================================
# DATABASE STORAGE
# =============================================================================

class InvestigationDatabase:
    """SQLite database for storing investigation data."""
    
    def __init__(self, db_path: Path):
        self.db_path = db_path
        self.conn = sqlite3.connect(str(db_path))
        self._closed = False
        self._init_schema()
    
    @property
    def is_closed(self) -> bool:
        """Check if database connection is closed."""
        return self._closed
    
    def ensure_open(self) -> None:
        """Ensure database connection is open, reconnect if needed."""
        if self._closed:
            self.conn = sqlite3.connect(str(self.db_path))
            self._closed = False
    
    def _init_schema(self) -> None:
        """Initialize database schema."""
        cursor = self.conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS investigations (
                id TEXT PRIMARY KEY,
                name TEXT NOT NULL,
                type TEXT NOT NULL,
                config TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS findings (
                id TEXT PRIMARY KEY,
                investigation_id TEXT NOT NULL,
                severity TEXT NOT NULL,
                title TEXT NOT NULL,
                description TEXT,
                evidence TEXT,
                recommendations TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (investigation_id) REFERENCES investigations(id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS entities (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                investigation_id TEXT NOT NULL,
                type TEXT NOT NULL,
                value TEXT NOT NULL,
                confidence REAL,
                source TEXT,
                context TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (investigation_id) REFERENCES investigations(id)
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS collected_data (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                investigation_id TEXT NOT NULL,
                data_type TEXT NOT NULL,
                data TEXT NOT NULL,
                source TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (investigation_id) REFERENCES investigations(id)
            )
        ''')
        
        self.conn.commit()
    
    def save_investigation(self, investigation: 'Investigation') -> None:
        """Save investigation to database."""
        cursor = self.conn.cursor()
        
        # Create JSON-serializable config dict
        config_dict = {
            'name': investigation.config.name,
            'investigation_type': investigation.config.investigation_type.name,
            'targets': investigation.config.targets,
            'data_sources': [str(p) for p in investigation.config.data_sources],
            'output_dir': str(investigation.config.output_dir),
            'api_keys': investigation.config.api_keys,
            'custom_patterns': investigation.config.custom_patterns,
            'max_depth': investigation.config.max_depth,
            'timeout': investigation.config.timeout,
            'parallel_workers': investigation.config.parallel_workers,
            'enable_caching': investigation.config.enable_caching,
            'cache_ttl': investigation.config.cache_ttl,
            'report_format': investigation.config.report_format
        }
        
        # Save investigation metadata
        cursor.execute('''
            INSERT OR REPLACE INTO investigations (id, name, type, config, updated_at)
            VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
        ''', (
            investigation.id,
            investigation.config.name,
            investigation.config.investigation_type.name,
            json.dumps(config_dict)
        ))
        
        # Save findings
        for finding in investigation.findings:
            cursor.execute('''
                INSERT OR REPLACE INTO findings 
                (id, investigation_id, severity, title, description, evidence, recommendations)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                finding.id,
                investigation.id,
                finding.severity.name,
                finding.title,
                finding.description,
                json.dumps(finding.evidence),
                json.dumps(finding.recommendations)
            ))
        
        self.conn.commit()
    
    def load_investigation(self, investigation_id: str) -> Optional[Dict[str, Any]]:
        """Load investigation from database."""
        cursor = self.conn.cursor()
        
        cursor.execute(
            'SELECT * FROM investigations WHERE id = ?',
            (investigation_id,)
        )
        row = cursor.fetchone()
        
        if row:
            return {
                'id': row[0],
                'name': row[1],
                'type': row[2],
                'config': json.loads(row[3]),
                'created_at': row[4],
                'updated_at': row[5]
            }
        
        return None
    
    def close(self) -> None:
        """Close database connection."""
        if not self._closed:
            self.conn.close()
            self._closed = True


# =============================================================================
# FILE WATCHDOG MONITORING SYSTEM
# =============================================================================

class InvestigationFileHandler:
    """Handles file system events for investigation monitoring.
    
    This handler detects new files added to monitored source directories
    and queues them for incremental analysis.
    """
    
    def __init__(self, investigation: 'Investigation', callback: Optional[Callable] = None):
        self.investigation = investigation
        self.callback = callback
        self.pending_files: List[Path] = []
        self.processed_files: Set[str] = set()
        self._lock = threading.Lock()
        self._load_processed_files()
    
    def _load_processed_files(self) -> None:
        """Load list of already processed files from investigation."""
        processed_file = self.investigation.config.output_dir / '.processed_files.json'
        if processed_file.exists():
            try:
                with processed_file.open('r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.processed_files = set(data.get('files', []))
            except Exception as e:
                logger.warning(f"Could not load processed files list: {e}")
    
    def _save_processed_files(self) -> None:
        """Save list of processed files."""
        processed_file = self.investigation.config.output_dir / '.processed_files.json'
        try:
            with processed_file.open('w', encoding='utf-8') as f:
                json.dump({'files': list(self.processed_files), 
                          'last_updated': format_timestamp()}, f, indent=2)
        except Exception as e:
            logger.warning(f"Could not save processed files list: {e}")
    
    def on_created(self, event) -> None:
        """Handle file creation event."""
        if hasattr(event, 'is_directory') and event.is_directory:
            return
        
        file_path = Path(event.src_path)
        self._queue_file(file_path)
    
    def on_modified(self, event) -> None:
        """Handle file modification event."""
        if hasattr(event, 'is_directory') and event.is_directory:
            return
        
        file_path = Path(event.src_path)
        # Only re-process if significantly modified
        file_key = str(file_path.absolute())
        if file_key in self.processed_files:
            # Check if file was truly modified (not just accessed)
            logger.debug(f"File already processed, skipping: {file_path}")
            return
        
        self._queue_file(file_path)
    
    def _queue_file(self, file_path: Path) -> None:
        """Add file to processing queue."""
        file_key = str(file_path.absolute())
        
        with self._lock:
            if file_key not in self.processed_files and file_path not in self.pending_files:
                # Ignore hidden files and report files
                if file_path.name.startswith('.'):
                    return
                if file_path.suffix in ['.db', '.db-journal']:
                    return
                if 'report' in file_path.name.lower():
                    return
                
                self.pending_files.append(file_path)
                logger.info(f"📁 New file detected: {file_path.name}")
                
                if self.callback:
                    self.callback(file_path)
    
    def mark_processed(self, file_path: Path) -> None:
        """Mark a file as processed."""
        file_key = str(file_path.absolute())
        with self._lock:
            self.processed_files.add(file_key)
            if file_path in self.pending_files:
                self.pending_files.remove(file_path)
        self._save_processed_files()
    
    def get_pending_files(self) -> List[Path]:
        """Get list of pending files."""
        with self._lock:
            return list(self.pending_files)
    
    def clear_pending(self) -> None:
        """Clear pending files list."""
        with self._lock:
            self.pending_files.clear()


class WatchdogEventHandler:
    """Watchdog event handler wrapper that works with or without watchdog installed."""
    
    def __init__(self, file_handler: InvestigationFileHandler):
        self.file_handler = file_handler
        
        if WATCHDOG_AVAILABLE:
            # Create actual watchdog handler
            class _Handler(FileSystemEventHandler):
                def __init__(self):
                    super().__init__()
                    self.handler = file_handler
                
                def on_created(self, event):
                    self.handler.on_created(event)
                
                def on_modified(self, event):
                    self.handler.on_modified(event)
            
            self._handler = _Handler()
        else:
            self._handler = None
    
    @property
    def handler(self):
        return self._handler


class InvestigationWatcher:
    """Watches source directories for new files and triggers incremental analysis.
    
    This class monitors all data sources in an investigation for new files.
    When new files are detected, they are queued for analysis and the investigation
    is updated with new findings.
    """
    
    def __init__(self, investigation: 'Investigation'):
        self.investigation = investigation
        self.file_handler = InvestigationFileHandler(
            investigation, 
            callback=self._on_new_file
        )
        self._observer = None
        self._running = False
        self._watch_thread = None
        self._new_file_event = threading.Event()
        self._auto_analyze = True
        self._analysis_delay = 5  # Seconds to wait before auto-analyzing
    
    def _on_new_file(self, file_path: Path) -> None:
        """Callback when new file is detected."""
        self._new_file_event.set()
    
    def start(self) -> bool:
        """Start watching source directories."""
        if not WATCHDOG_AVAILABLE:
            logger.warning("Watchdog not installed. Install with: pip install watchdog")
            return False
        
        if self._running:
            logger.warning("Watcher already running")
            return True
        
        try:
            self._observer = Observer()
            event_handler = WatchdogEventHandler(self.file_handler)
            
            # Watch all data source directories
            watched_dirs = set()
            for source in self.investigation.config.data_sources:
                if source.is_dir():
                    watch_path = source
                else:
                    watch_path = source.parent
                
                if watch_path.exists() and str(watch_path) not in watched_dirs:
                    self._observer.schedule(
                        event_handler.handler, 
                        str(watch_path), 
                        recursive=True
                    )
                    watched_dirs.add(str(watch_path))
                    logger.info(f"👁️ Watching: {watch_path}")
            
            if not watched_dirs:
                logger.warning("No valid directories to watch")
                return False
            
            self._observer.start()
            self._running = True
            
            # Start background thread for auto-analysis
            if self._auto_analyze:
                self._watch_thread = threading.Thread(
                    target=self._auto_analyze_loop, 
                    daemon=True
                )
                self._watch_thread.start()
            
            logger.info("🔍 File watcher started")
            return True
            
        except Exception as e:
            logger.error(f"Failed to start watcher: {e}")
            return False
    
    def stop(self) -> None:
        """Stop watching directories."""
        self._running = False
        self._new_file_event.set()  # Wake up the thread
        
        if self._observer:
            self._observer.stop()
            self._observer.join(timeout=5)
            self._observer = None
        
        logger.info("🛑 File watcher stopped")
    
    def _auto_analyze_loop(self) -> None:
        """Background loop that triggers analysis when new files are detected."""
        while self._running:
            # Wait for new file event or timeout
            triggered = self._new_file_event.wait(timeout=self._analysis_delay)
            
            if not self._running:
                break
            
            if triggered:
                self._new_file_event.clear()
                # Wait a bit more for additional files
                import time
                time.sleep(2)
                
                # Process pending files
                pending = self.file_handler.get_pending_files()
                if pending:
                    logger.info(f"📊 Auto-analyzing {len(pending)} new file(s)...")
                    self.analyze_new_files()
    
    def analyze_new_files(self) -> int:
        """Analyze newly detected files and update investigation.
        
        Returns:
            Number of new files analyzed.
        """
        pending = self.file_handler.get_pending_files()
        if not pending:
            logger.info("No new files to analyze")
            return 0
        
        logger.info(f"Analyzing {len(pending)} new file(s)...")
        
        # Collect data from new files
        collector = FileCollector()
        new_data = []
        
        for file_path in pending:
            try:
                for record in collector.collect(str(file_path), self.investigation.config):
                    new_data.append(record)
                self.file_handler.mark_processed(file_path)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
        
        if new_data:
            # Add to investigation data
            self.investigation.collected_data.extend(new_data)
            logger.info(f"Collected {len(new_data)} new records")
            
            # Run incremental analysis
            self.investigation.run_incremental(new_data)
        
        return len(pending)
    
    @property
    def is_running(self) -> bool:
        return self._running
    
    @property
    def pending_count(self) -> int:
        return len(self.file_handler.get_pending_files())


# =============================================================================
# MAIN INVESTIGATION CLASS
# =============================================================================

@dataclass
class Investigation:
    """Main investigation class that orchestrates the entire process."""
    
    config: InvestigationConfig
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    collected_data: List[Dict[str, Any]] = field(default_factory=list)
    findings: List[Finding] = field(default_factory=list)
    entities: List[Entity] = field(default_factory=list)
    timeline: List[TimelineEvent] = field(default_factory=list)
    _collectors: List[DataCollector] = field(default_factory=list)
    _analyzers: List[AnalysisModule] = field(default_factory=list)
    _cache: Optional[Cache] = field(default=None)
    _db: Optional[InvestigationDatabase] = field(default=None)
    _watcher: Optional[InvestigationWatcher] = field(default=None)
    _previous_analysis: Optional[Dict[str, Any]] = field(default=None)
    _run_count: int = field(default=0)
    
    def __post_init__(self):
        """Initialize investigation components."""
        # Set up collectors
        self._collectors = [
            FileCollector(),
            WeatherCollector(),
            PublicRecordsCollector(),
        ]
        
        # Set up analyzers
        self._analyzers = [
            EntityExtractionModule(),
            TimelineReconstructionModule(),
            AnomalyDetectionModule(),
            RelationshipMappingModule(),
            LLMEnhancedAnalysisModule(),  # Ollama LLM integration
        ]
        
        # Set up cache
        if self.config.enable_caching:
            cache_dir = self.config.output_dir / '.cache'
            self._cache = Cache(cache_dir, self.config.cache_ttl)
        
        # Set up database
        db_path = self.config.output_dir / 'investigation.db'
        self.config.output_dir.mkdir(parents=True, exist_ok=True)
        self._db = InvestigationDatabase(db_path)
        
        # Load previous analysis if exists
        self._load_previous_analysis()
    
    def _load_previous_analysis(self) -> None:
        """Load previous analysis data from existing reports."""
        # Look for existing report.json in output directory
        report_path = self.config.output_dir / 'report.json'
        
        if report_path.exists():
            try:
                with report_path.open('r', encoding='utf-8') as f:
                    self._previous_analysis = json.load(f)
                    
                # Restore previous findings count for context
                if self._previous_analysis:
                    prev_findings = self._previous_analysis.get('summary', {}).get('total_findings', 0)
                    prev_records = self._previous_analysis.get('summary', {}).get('total_records', 0)
                    self._run_count = self._previous_analysis.get('metadata', {}).get('run_count', 0)
                    
                    logger.info(f"📚 Loaded previous analysis: {prev_findings} findings from {prev_records} records (run #{self._run_count})")
                
            except Exception as e:
                logger.warning(f"Could not load previous analysis: {e}")
                self._previous_analysis = None
    
    def _get_previous_context(self) -> str:
        """Get a summary of previous analysis for LLM context."""
        if not self._previous_analysis:
            return ""
        
        context_parts = []
        
        # Add previous investigation summary
        inv_data = self._previous_analysis.get('investigation', {})
        context_parts.append(f"Previous Analysis (Run #{self._run_count}):")
        context_parts.append(f"- Investigation: {inv_data.get('name', 'Unknown')}")
        context_parts.append(f"- Timestamp: {inv_data.get('timestamp', 'Unknown')}")
        
        # Add previous findings summary
        prev_findings = self._previous_analysis.get('findings', [])
        if prev_findings:
            context_parts.append(f"\nPrevious Findings ({len(prev_findings)} total):")
            for finding in prev_findings[:10]:  # Limit to recent findings
                context_parts.append(f"- [{finding.get('severity', 'INFO')}] {finding.get('title', 'Unknown')}")
        
        # Add previous entity summary if available
        if 'entities_summary' in self._previous_analysis:
            entities = self._previous_analysis['entities_summary']
            context_parts.append(f"\nPreviously Identified Entities: {entities}")
        
        return "\n".join(context_parts)
    
    def run(self) -> None:
        """Run the complete investigation."""
        self._run_count += 1
        logger.info(f"Starting investigation: {self.config.name} (Run #{self._run_count})")
        
        try:
            # Phase 1: Data Collection
            logger.info("Phase 1: Collecting data...")
            self._collect_data()
            
            # Phase 2: Analysis (with previous context)
            logger.info("Phase 2: Analyzing data...")
            self._analyze_data()
            
            # Phase 3: Report Generation (single report files, overwritten each time)
            logger.info("Phase 3: Generating reports...")
            self._generate_reports()
            
            # Phase 4: Save to database
            logger.info("Phase 4: Saving to database...")
            if self._db:
                self._db.ensure_open()
                self._db.save_investigation(self)
            
            logger.info(f"Investigation complete. Found {len(self.findings)} findings.")
            
        except Exception as e:
            logger.error(f"Investigation failed: {e}")
            raise
    
    def run_incremental(self, new_data: List[Dict[str, Any]]) -> None:
        """Run incremental analysis on new data only.
        
        This method is called by the file watcher when new files are detected.
        It analyzes only the new data and merges findings with existing ones.
        
        Args:
            new_data: New data records to analyze.
        """
        self._run_count += 1
        logger.info(f"Running incremental analysis (Run #{self._run_count})...")
        
        try:
            # Get previous context for LLM
            prev_context = self._get_previous_context()
            
            # Analyze new data
            new_findings = []
            for analyzer in self._analyzers:
                # For LLM analyzer, pass previous context
                if isinstance(analyzer, LLMEnhancedAnalysisModule):
                    # The LLM will have access to previous findings through the investigation
                    findings = analyzer.analyze(new_data, self.config)
                else:
                    findings = analyzer.analyze(new_data, self.config)
                new_findings.extend(findings)
            
            if new_findings:
                # Add new findings with incremental marker
                # Since Finding is a NamedTuple (immutable), create new instances
                marked_findings = []
                for finding in new_findings:
                    marked_finding = Finding(
                        id=finding.id,
                        severity=finding.severity,
                        title=finding.title,
                        description=f"[Incremental Run #{self._run_count}] {finding.description}",
                        entities=finding.entities,
                        evidence=finding.evidence,
                        timestamp=finding.timestamp,
                        recommendations=finding.recommendations
                    )
                    marked_findings.append(marked_finding)
                
                self.findings.extend(marked_findings)
                logger.info(f"Generated {len(marked_findings)} new findings from incremental analysis")
            
            # Regenerate reports with all findings
            self._generate_reports()
            
            # Update database
            if self._db:
                self._db.ensure_open()
                self._db.save_investigation(self)
            
            logger.info(f"Incremental analysis complete. Total findings: {len(self.findings)}")
            
        except Exception as e:
            logger.error(f"Incremental analysis failed: {e}")
    
    def start_watching(self) -> bool:
        """Start file watcher for source directories.
        
        Returns:
            True if watcher started successfully, False otherwise.
        """
        if not WATCHDOG_AVAILABLE:
            logger.error("Watchdog not available. Install with: pip install watchdog")
            return False
        
        if self._watcher is None:
            self._watcher = InvestigationWatcher(self)
        
        return self._watcher.start()
    
    def stop_watching(self) -> None:
        """Stop file watcher."""
        if self._watcher:
            self._watcher.stop()
    
    @property
    def is_watching(self) -> bool:
        """Check if file watcher is running."""
        return self._watcher is not None and self._watcher.is_running
    
    def _collect_data(self) -> None:
        """Collect data from all sources."""
        # Collect from file sources
        for source in self.config.data_sources:
            if source.exists():
                for collector in self._collectors:
                    if isinstance(collector, FileCollector):
                        for record in collector.collect(str(source), self.config):
                            self.collected_data.append(record)
        
        # Collect from targets
        for target in self.config.targets:
            for collector in self._collectors:
                if collector.supports_type(self.config.investigation_type):
                    # Check cache first
                    cache_key = f"{collector.__class__.__name__}:{target}"
                    if self._cache:
                        cached = self._cache.get(cache_key)
                        if cached:
                            self.collected_data.extend(cached.get('data', []))
                            continue
                    
                    # Collect fresh data
                    records = list(collector.collect(target, self.config))
                    self.collected_data.extend(records)
                    
                    # Cache results
                    if self._cache:
                        self._cache.set(cache_key, {'data': records})
        
        logger.info(f"Collected {len(self.collected_data)} data records")
    
    def _analyze_data(self) -> None:
        """Analyze collected data."""
        for analyzer in self._analyzers:
            findings = analyzer.analyze(self.collected_data, self.config)
            self.findings.extend(findings)
        
        logger.info(f"Generated {len(self.findings)} findings")
    
    def _generate_reports(self) -> None:
        """Generate investigation reports.
        
        Always generates report.md and report.json in the output directory.
        These files are overwritten on each run, building on previous analysis.
        """
        # Always use fixed filenames - one MD and one JSON
        md_path = self.config.output_dir / 'report.md'
        json_path = self.config.output_dir / 'report.json'
        
        # Generate Markdown report
        MarkdownReportGenerator().generate(self, md_path)
        logger.info(f"Report generated: {md_path}")
        
        # Generate JSON report with metadata for future runs
        self._generate_json_report_with_metadata(json_path)
        logger.info(f"Report generated: {json_path}")
    
    def _generate_json_report_with_metadata(self, output_path: Path) -> None:
        """Generate JSON report with metadata for incremental analysis."""
        # Build entities summary for future context
        entities_summary = {}
        for finding in self.findings:
            for entity in finding.entities:
                entity_type = entity.type.name
                if entity_type not in entities_summary:
                    entities_summary[entity_type] = []
                if entity.value not in entities_summary[entity_type]:
                    entities_summary[entity_type].append(entity.value)
        
        report = {
            'metadata': {
                'run_count': self._run_count,
                'last_updated': format_timestamp(),
                'investigation_id': self.id,
            },
            'investigation': {
                'name': self.config.name,
                'type': self.config.investigation_type.name,
                'targets': self.config.targets,
                'timestamp': format_timestamp()
            },
            'summary': {
                'total_records': len(self.collected_data),
                'total_findings': len(self.findings),
                'severity_distribution': {
                    severity.name: sum(
                        1 for f in self.findings if f.severity == severity
                    )
                    for severity in SeverityLevel
                }
            },
            'entities_summary': entities_summary,
            'findings': [
                {
                    'id': f.id,
                    'severity': f.severity.name,
                    'title': f.title,
                    'description': f.description,
                    'evidence': f.evidence,
                    'recommendations': f.recommendations,
                    'timestamp': f.timestamp.isoformat() if hasattr(f.timestamp, 'isoformat') else str(f.timestamp)
                }
                for f in self.findings
            ],
            'data_sources': [str(s) for s in self.config.data_sources]
        }
        
        with output_path.open('w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)


# =============================================================================
# COMMAND LINE INTERFACE
# =============================================================================

class Command(ABC):
    """Abstract base class for CLI commands."""
    
    @abstractmethod
    def execute(self, options: argparse.Namespace) -> None:
        """Execute the command."""
        pass


class InitCommand(Command):
    """Initialize a new investigation."""
    
    def execute(self, options: argparse.Namespace) -> None:
        """Create new investigation configuration."""
        # Create investigation directory inside INVESTIGATIONS folder
        inv_dir_name = options.name.replace(' ', '_').lower()
        output_dir = BASE_INVESTIGATIONS_DIR / inv_dir_name
        
        config = InvestigationConfig(
            name=options.name,
            investigation_type=InvestigationType[options.type.upper()],
            targets=options.targets.split(',') if options.targets else [],
            data_sources=[Path(p) for p in options.sources.split(',')] if options.sources else [],
            output_dir=output_dir,
        )
        
        config_path = output_dir / 'config.json'
        output_dir.mkdir(parents=True, exist_ok=True)
        config.to_file(config_path)
        
        print(f"Investigation initialized: {config_path}")
        print(f"Output directory: {output_dir}")


class RunCommand(Command):
    """Run an investigation."""
    
    def execute(self, options: argparse.Namespace) -> None:
        """Run investigation from configuration."""
        config_path = Path(options.config)
        config = InvestigationConfig.from_file(config_path)
        
        investigation = Investigation(config=config)
        investigation.run()
        
        print(f"Investigation complete. Check {config.output_dir} for results.")


class AnalyzeCommand(Command):
    """Analyze a specific data source."""
    
    def execute(self, options: argparse.Namespace) -> None:
        """Quick analysis of a data source."""
        source_path = Path(options.source)
        
        # Create analysis directory inside INVESTIGATIONS folder
        analysis_name = f"quick_analysis_{source_path.stem}_{format_timestamp_compact()}"
        output_dir = BASE_INVESTIGATIONS_DIR / analysis_name
        
        config = InvestigationConfig(
            name=f"Quick Analysis: {source_path.name}",
            investigation_type=InvestigationType.INCIDENT,
            targets=[],
            data_sources=[source_path],
            output_dir=output_dir,
        )
        
        investigation = Investigation(config=config)
        investigation.run()
        print(f"Analysis output saved to: {output_dir}")


class ExtractCommand(Command):
    """Extract entities from a file."""
    
    def execute(self, options: argparse.Namespace) -> None:
        """Extract and display entities."""
        source_path = Path(options.source)
        
        with source_path.open() as f:
            content = f.read()
        
        entities = PatternLibrary.extract_entities(content, source=str(source_path))
        
        if options.format == 'json':
            output = [
                {
                    'type': e.type.name,
                    'value': e.value,
                    'confidence': e.confidence
                }
                for e in entities
            ]
            print(json.dumps(output, indent=2))
        else:
            for entity in entities:
                print(f"{entity.type.name}: {entity.value}")


class SearchCommand(Command):
    """Search across investigation data."""
    
    def execute(self, options: argparse.Namespace) -> None:
        """Search for patterns in data."""
        db_path = Path(options.database)
        
        if not db_path.exists():
            print(f"Database not found: {db_path}")
            return
        
        db = InvestigationDatabase(db_path)
        cursor = db.conn.cursor()
        
        # Search entities
        cursor.execute('''
            SELECT type, value, source FROM entities
            WHERE value LIKE ?
        ''', (f'%{options.query}%',))
        
        results = cursor.fetchall()
        
        for row in results:
            print(f"{row[0]}: {row[1]} (source: {row[2]})")
        
        db.close()


def create_parser() -> argparse.ArgumentParser:
    """Create command-line argument parser."""
    parser = argparse.ArgumentParser(
        prog='osint-framework',
        description='OSINT Investigation Framework - Comprehensive intelligence gathering and analysis',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f'''
Examples:
  Initialize new investigation:
    osint-framework init --name "Case 001" --type person

  Run investigation:
    osint-framework run --config ./INVESTIGATIONS/case_001/config.json

  Quick analysis:
    osint-framework analyze --source ./logs/access.log

  Extract entities:
    osint-framework extract --source ./document.txt --format json

  Search database:
    osint-framework search --database ./INVESTIGATIONS/case_001/investigation.db --query "192.168"

Note: All investigation output is saved to: {BASE_INVESTIGATIONS_DIR}
        '''
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Available commands')
    
    # Init command
    init_parser = subparsers.add_parser('init', help='Initialize new investigation')
    init_parser.add_argument('--name', '-n', required=True, help='Investigation name (used as folder name)')
    init_parser.add_argument(
        '--type', '-t', 
        choices=[t.name.lower() for t in InvestigationType],
        default='incident',
        help='Investigation type'
    )
    init_parser.add_argument('--targets', help='Comma-separated list of targets')
    init_parser.add_argument('--sources', help='Comma-separated list of data sources')
    
    # Run command
    run_parser = subparsers.add_parser('run', help='Run investigation')
    run_parser.add_argument('--config', '-c', required=True, help='Configuration file path')
    
    # Analyze command
    analyze_parser = subparsers.add_parser('analyze', help='Quick analysis of data source')
    analyze_parser.add_argument('--source', '-s', required=True, help='Data source path')
    
    # Extract command
    extract_parser = subparsers.add_parser('extract', help='Extract entities from file')
    extract_parser.add_argument('--source', '-s', required=True, help='Source file path')
    extract_parser.add_argument('--format', '-f', choices=['text', 'json'], default='text')
    
    # Search command
    search_parser = subparsers.add_parser('search', help='Search investigation database')
    search_parser.add_argument('--database', '-d', required=True, help='Database path')
    search_parser.add_argument('--query', '-q', required=True, help='Search query')
    
    return parser


def main(argv: List[str] = sys.argv[1:]) -> int:
    """Main entry point."""
    parser = create_parser()
    options = parser.parse_args(argv)
    
    if not options.command:
        parser.print_help()
        return 1
    
    commands = {
        'init': InitCommand(),
        'run': RunCommand(),
        'analyze': AnalyzeCommand(),
        'extract': ExtractCommand(),
        'search': SearchCommand(),
    }
    
    try:
        command = commands.get(options.command)
        if command:
            command.execute(options)
            return 0
        else:
            print(f"Unknown command: {options.command}")
            return 1
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1
    finally:
        logging.shutdown()


# =============================================================================
# INTERACTIVE MODE
# =============================================================================

class InteractiveSession:
    """Interactive investigation session."""
    
    def __init__(self):
        self.investigation: Optional[Investigation] = None
        self.running = True
        self.current_model = "wizardlm2:latest"  # Default LLM model optimized for GTX 1070
        self.ollama_process: Optional[subprocess.Popen] = None  # Track Ollama process
    
    def start(self) -> None:
        """Start interactive session."""
        # Ensure Ollama server is running
        print("Checking Ollama server status...")
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            if response.status_code == 200:
                print("✓ Ollama server is already running.")
            else:
                raise Exception("Server not responding")
        except Exception as e:
            print(f"⚠ Ollama server not running ({e}). Starting server...")
            try:
                import subprocess
                self.ollama_process = subprocess.Popen([
                    r"C:\Users\davee\AppData\Local\Programs\Ollama\ollama.exe",
                    "serve"
                ])
                print("Waiting 10 seconds for Ollama to initialize...")
                import time
                time.sleep(10)
                print("✓ Ollama server started.")
            except Exception as start_error:
                print(f"✗ Failed to start Ollama server: {start_error}")
                print("Continuing without LLM analysis capabilities.")
        
        print("\n" + "="*60)
        print("OSINT Investigation Framework - Interactive Mode")
        print("="*60)
        print(f"\nInvestigations Directory: {BASE_INVESTIGATIONS_DIR}")
        print("\nType 'help' for available commands.\n")
        
        while self.running:
            try:
                command = input("osint> ").strip()
                self._process_command(command)
            except KeyboardInterrupt:
                print("\nInterrupted. Cleaning up...")
                self._cmd_exit([])
                break
            except EOFError:
                self._cmd_exit([])
                break
    
    def _process_command(self, command: str) -> None:
        """Process interactive command."""
        parts = command.split()
        if not parts:
            return
        
        cmd = parts[0].lower()
        args = parts[1:]
        
        handlers = {
            'help': self._cmd_help,
            'new': self._cmd_new,
            'load': self._cmd_load,
            'add': self._cmd_add,
            'run': self._cmd_run,
            'watch': self._cmd_watch,
            'unwatch': self._cmd_unwatch,
            'findings': self._cmd_findings,
            'entities': self._cmd_entities,
            'export': self._cmd_export,
            'status': self._cmd_status,
            'model': self._cmd_model,
            'models': self._cmd_models,
            'clear': self._cmd_clear,
            'cls': self._cmd_clear,
            'exit': self._cmd_exit,
            'quit': self._cmd_exit,
        }
        
        handler = handlers.get(cmd, self._cmd_unknown)
        handler(args)
    
    def _cmd_help(self, args: List[str]) -> None:
        """Display help."""
        watch_status = ""
        if self.investigation and hasattr(self.investigation, 'is_watching') and self.investigation.is_watching:
            watch_status = " [WATCHING]"
        
        print(f"""
Available commands:
  help              - Show this help message
  new <type> <name> - Create new investigation of specified type
  load <name>       - Load investigation by name (or full config path)
  load              - List available investigations
  add target <t>    - Add investigation target
  add source <s>    - Add data source (directory or file)
  run               - Execute investigation
  watch             - Start watching sources for new files{watch_status}
  unwatch           - Stop watching for new files
  findings          - Display findings
  entities          - Display extracted entities
  export <format>   - Export report (markdown/json/csv)
  status            - Show investigation status
  models            - List available Ollama models
  model <name>      - Switch LLM model (current: {self.current_model})
  clear/cls         - Clear the terminal screen
  exit/quit         - Exit interactive mode

Investigation Types:
  PERSON, ORGANIZATION, DOMAIN, IP_ADDRESS, EMAIL, PHONE, SOCIAL_MEDIA,
  CRYPTOCURRENCY, VEHICLE, LOCATION, INCIDENT, NETWORK, MALWARE

Examples:
  new PERSON "John Doe"     - Create person investigation
  new DOMAIN example.com    - Create domain investigation
  load nunley              - Load the 'nunley' investigation
  load "my case"           - Load investigation with spaces in name

Watchdog Mode:
  When 'watch' is active, the framework monitors source directories
  for new files. New files are automatically analyzed and findings
  are merged into the existing report (report.md / report.json).
  The LLM builds on previous analysis for continuity.
        """)
    
    def _cmd_clear(self, args: List[str]) -> None:
        """Clear the terminal screen."""
        import os
        os.system('cls' if os.name == 'nt' else 'clear')
    
    def _cmd_new(self, args: List[str]) -> None:
        """Create new investigation."""
        if not args:
            print("Usage: new <type> <name>")
            print("Types: PERSON, ORGANIZATION, DOMAIN, IP_ADDRESS, EMAIL, PHONE, SOCIAL_MEDIA, CRYPTOCURRENCY, VEHICLE, LOCATION, INCIDENT, NETWORK, MALWARE")
            return
        
        inv_type_str = args[0].upper()
        try:
            inv_type = InvestigationType[inv_type_str]
        except KeyError:
            print(f"Unknown investigation type: {inv_type_str}")
            print("Valid types: PERSON, ORGANIZATION, DOMAIN, IP_ADDRESS, EMAIL, PHONE, SOCIAL_MEDIA, CRYPTOCURRENCY, VEHICLE, LOCATION, INCIDENT, NETWORK, MALWARE")
            return
        
        name = ' '.join(args[1:]) if len(args) > 1 else input("Investigation name: ")
        
        # Create investigation directory inside INVESTIGATIONS folder
        inv_dir_name = name.replace(' ', '_').lower()
        output_dir = BASE_INVESTIGATIONS_DIR / inv_dir_name
        output_dir.mkdir(parents=True, exist_ok=True)
        
        config = InvestigationConfig(
            name=name,
            investigation_type=inv_type,
            targets=[],
            data_sources=[],
            output_dir=output_dir,
        )
        
        # Save config.json for later loading
        config_path = output_dir / 'config.json'
        config.to_file(config_path)
        
        self.investigation = Investigation(config=config)
        print(f"Created {inv_type.value} investigation: {name}")
        print(f"Output directory: {output_dir}")
        print(f"Config saved to: {config_path}")
    
    def _cmd_load(self, args: List[str]) -> None:
        """Load investigation by name or config path.
        
        Usage:
            load nunley              - Load from INVESTIGATIONS/nunley/config.json
            load "my case"           - Load from INVESTIGATIONS/my case/config.json
            load C:/path/config.json - Load from full path
        """
        if not args:
            # Show available investigations
            print("\nAvailable investigations:")
            if BASE_INVESTIGATIONS_DIR.exists():
                investigations = [d.name for d in BASE_INVESTIGATIONS_DIR.iterdir() 
                                if d.is_dir() and (d / "config.json").exists()]
                if investigations:
                    for inv in sorted(investigations):
                        print(f"  - {inv}")
                    print(f"\nUsage: load <name>")
                else:
                    print("  No investigations found.")
            return
        
        # Join args to handle names with spaces
        name_or_path = ' '.join(args).strip('"').strip("'")
        
        # Check if it's a direct path to config file
        direct_path = Path(name_or_path)
        if direct_path.exists() and direct_path.is_file():
            config_path = direct_path
        else:
            # Treat as investigation name - look in INVESTIGATIONS directory
            investigation_dir = BASE_INVESTIGATIONS_DIR / name_or_path
            config_path = investigation_dir / "config.json"
            
            if not config_path.exists():
                # Try case-insensitive search
                found = False
                if BASE_INVESTIGATIONS_DIR.exists():
                    for d in BASE_INVESTIGATIONS_DIR.iterdir():
                        if d.is_dir() and d.name.lower() == name_or_path.lower():
                            config_path = d / "config.json"
                            if config_path.exists():
                                found = True
                                break
                
                if not found:
                    print(f"Investigation not found: {name_or_path}")
                    print(f"Looked in: {investigation_dir}")
                    print("\nAvailable investigations:")
                    if BASE_INVESTIGATIONS_DIR.exists():
                        investigations = [d.name for d in BASE_INVESTIGATIONS_DIR.iterdir() 
                                        if d.is_dir() and (d / "config.json").exists()]
                        for inv in sorted(investigations):
                            print(f"  - {inv}")
                    return
        
        if config_path.exists():
            config = InvestigationConfig.from_file(config_path)
            self.investigation = Investigation(config=config)
            print(f"Loaded investigation: {config.name}")
            print(f"Output directory: {config.output_dir}")
            print(f"Targets: {config.targets}")
            print(f"Data sources: {[str(s) for s in config.data_sources]}")
        else:
            print(f"Config not found: {config_path}")
    
    def _cmd_add(self, args: List[str]) -> None:
        """Add target or source."""
        if not self.investigation:
            print("No active investigation. Use 'new' first.")
            return
        
        if len(args) < 2:
            print("Usage: add target <value> or add source <path>")
            return
        
        add_type = args[0].lower()
        value = ' '.join(args[1:])  # Allow spaces in paths/values
        # Strip any surrounding quotes that user might have included
        value = value.strip('"').strip("'")
        
        if add_type == 'target':
            self.investigation.config.targets.append(value)
            print(f"Added target: {value}")
        elif add_type == 'source':
            source_path = Path(value)
            if source_path.exists():
                self.investigation.config.data_sources.append(source_path)
                print(f"Added source: {value}")
            else:
                print(f"Warning: Path does not exist: {value}")
                confirm = input("Add anyway? (y/n): ").strip().lower()
                if confirm == 'y':
                    self.investigation.config.data_sources.append(source_path)
                    print(f"Added source: {value}")
                else:
                    print("Source not added.")
                    return
        else:
            print(f"Unknown add type: {add_type}")
            return
        
        # Auto-save config after changes
        config_path = self.investigation.config.output_dir / 'config.json'
        self.investigation.config.to_file(config_path)
        print(f"Config updated: {config_path}")
    
    def _cmd_run(self, args: List[str]) -> None:
        """Run investigation."""
        if not self.investigation:
            print("No active investigation. Use 'new' first.")
            return
        
        self.investigation.run()
    
    def _cmd_watch(self, args: List[str]) -> None:
        """Start watching source directories for new files."""
        if not self.investigation:
            print("No active investigation. Use 'new' first.")
            return
        
        if not WATCHDOG_AVAILABLE:
            print("❌ Watchdog not installed.")
            print("   Install with: pip install watchdog")
            return
        
        if not self.investigation.config.data_sources:
            print("❌ No data sources configured. Use 'add source <path>' first.")
            return
        
        if self.investigation.is_watching:
            print("👁️ Already watching for new files.")
            self._show_watch_status()
            return
        
        print("🔍 Starting file watcher...")
        if self.investigation.start_watching():
            print("✅ Watcher started successfully!")
            print("\n📁 Monitoring directories:")
            for source in self.investigation.config.data_sources:
                watch_dir = source if source.is_dir() else source.parent
                print(f"   • {watch_dir}")
            print("\n💡 New files will be automatically analyzed.")
            print("   Use 'unwatch' to stop monitoring.")
            print("   Use 'status' to check pending files.")
        else:
            print("❌ Failed to start watcher. Check logs for details.")
    
    def _cmd_unwatch(self, args: List[str]) -> None:
        """Stop watching for new files."""
        if not self.investigation:
            print("No active investigation.")
            return
        
        if not self.investigation.is_watching:
            print("Not currently watching.")
            return
        
        self.investigation.stop_watching()
        print("🛑 File watcher stopped.")
    
    def _show_watch_status(self) -> None:
        """Display watch status details."""
        if self.investigation and self.investigation._watcher:
            pending = self.investigation._watcher.pending_count
            print(f"\n📊 Watch Status:")
            print(f"   Pending files: {pending}")
            if pending > 0:
                print("   Run 'run' to analyze pending files manually, or wait for auto-analysis.")
    
    def _cmd_findings(self, args: List[str]) -> None:
        """Display findings."""
        if not self.investigation:
            print("No active investigation.")
            return
        
        for finding in self.investigation.findings:
            print(f"\n[{finding.severity.name}] {finding.title}")
            print(f"  {finding.description}")
    
    def _cmd_entities(self, args: List[str]) -> None:
        """Display entities."""
        if not self.investigation:
            print("No active investigation.")
            return
        
        # Extract entities from findings
        for finding in self.investigation.findings:
            for entity in finding.entities:
                print(f"  {entity.type.name}: {entity.value}")
    
    def _cmd_export(self, args: List[str]) -> None:
        """Export report."""
        if not self.investigation:
            print("No active investigation.")
            return
        
        format_type = args[0] if args else 'markdown'
        self.investigation.config.report_format = format_type
        self.investigation._generate_reports()
        print(f"Report exported to {self.investigation.config.output_dir}")
    
    def _cmd_status(self, args: List[str]) -> None:
        """Show status."""
        if not self.investigation:
            print("No active investigation.")
            return
        
        print(f"\n{'='*50}")
        print(f"Investigation: {self.investigation.config.name}")
        print(f"{'='*50}")
        print(f"Type: {self.investigation.config.investigation_type.name}")
        print(f"Output Directory: {self.investigation.config.output_dir}")
        print(f"Run Count: {self.investigation._run_count}")
        print(f"\n📊 Data:")
        print(f"   Targets: {len(self.investigation.config.targets)}")
        print(f"   Sources: {len(self.investigation.config.data_sources)}")
        print(f"   Records collected: {len(self.investigation.collected_data)}")
        print(f"   Findings: {len(self.investigation.findings)}")
        
        # Show watch status
        if self.investigation.is_watching:
            print(f"\n👁️ File Watcher: ACTIVE")
            if self.investigation._watcher:
                pending = self.investigation._watcher.pending_count
                print(f"   Pending files: {pending}")
        else:
            print(f"\n👁️ File Watcher: Inactive")
            print(f"   Use 'watch' to start monitoring for new files")
        
        # Show report files
        report_md = self.investigation.config.output_dir / 'report.md'
        report_json = self.investigation.config.output_dir / 'report.json'
        print(f"\n📄 Reports:")
        if report_md.exists():
            print(f"   ✅ {report_md.name}")
        else:
            print(f"   ⬜ {report_md.name} (not generated)")
        if report_json.exists():
            print(f"   ✅ {report_json.name}")
        else:
            print(f"   ⬜ {report_json.name} (not generated)")
        print()
    
    def _cmd_models(self, args: List[str]) -> None:
        """List available Ollama models."""
        import requests
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=10)
            if response.status_code == 200:
                data = response.json()
                models = data.get("models", [])
                if models:
                    print("\n=== Available Ollama Models ===")
                    for model in models:
                        name = model.get("name", "unknown")
                        size = model.get("size", 0)
                        size_gb = size / (1024**3) if size else 0
                        modified = model.get("modified_at", "")[:10] if model.get("modified_at") else ""
                        print(f"  • {name} ({size_gb:.1f} GB) {modified}")
                    print(f"\nTotal: {len(models)} model(s)")
                    if hasattr(self, 'llm_module') and self.llm_module:
                        print(f"Current model: {self.llm_module.llm.model}")
                else:
                    print("No models found. Pull a model with: ollama pull <model_name>")
            else:
                print(f"Error: Could not fetch models (HTTP {response.status_code})")
        except requests.exceptions.ConnectionError:
            print("Error: Cannot connect to Ollama. Is it running?")
        except Exception as e:
            print(f"Error listing models: {e}")
    
    def _cmd_model(self, args: List[str]) -> None:
        """Switch the active Ollama model."""
        import requests
        if not args:
            if hasattr(self, 'llm_module') and self.llm_module:
                print(f"Current model: {self.llm_module.llm.model}")
            else:
                print("LLM module not initialized.")
            print("Usage: model <model_name>")
            print("Use 'models' to list available models.")
            return
        
        model_name = args[0]
        
        # Verify model exists
        try:
            response = requests.get("http://localhost:11434/api/tags", timeout=10)
            if response.status_code == 200:
                data = response.json()
                available = [m.get("name", "") for m in data.get("models", [])]
                if model_name not in available:
                    print(f"Warning: Model '{model_name}' not found in available models.")
                    print(f"Available: {', '.join(available)}")
                    confirm = input("Try to use it anyway? (y/n): ").strip().lower()
                    if confirm != 'y':
                        return
        except Exception:
            print("Warning: Could not verify model availability.")
        
        # Update the model
        if hasattr(self, 'llm_module') and self.llm_module:
            old_model = self.llm_module.llm.model
            self.llm_module.llm.model = model_name
            self.current_model = model_name  # Update display model
            print(f"Model changed: {old_model} -> {model_name}")
        else:
            print("LLM module not initialized. Initializing now...")
            self.llm_module = LLMEnhancedAnalysisModule(model=model_name)
            self.current_model = model_name  # Update display model
            print(f"LLM module initialized with model: {model_name}")
    
    def _cmd_exit(self, args: List[str]) -> None:
        """Exit session."""
        # Stop file watcher if running
        if self.investigation and self.investigation.is_watching:
            print("Stopping file watcher...")
            self.investigation.stop_watching()
        
        # Shut down Ollama if we started it
        if self.ollama_process:
            print("Shutting down Ollama server...")
            try:
                self.ollama_process.terminate()
                self.ollama_process.wait(timeout=5)
                print("✓ Ollama server shut down.")
            except Exception as e:
                print(f"⚠ Failed to shut down Ollama gracefully: {e}")
                try:
                    self.ollama_process.kill()
                except:
                    pass
        
        self.running = False
        print("Goodbye!")
    
    def _cmd_unknown(self, args: List[str]) -> None:
        """Unknown command."""
        print("Unknown command. Type 'help' for available commands.")


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    if len(sys.argv) > 1:
        # Command-line mode
        sys.exit(main())
    else:
        # Interactive mode
        session = InteractiveSession()
        session.start()